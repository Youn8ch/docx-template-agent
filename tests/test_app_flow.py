import json
import inspect
from pathlib import Path

import pytest
from docx import Document

import app
from src.engine.model.operation_model import ExecutionReport, SAFE_OPERATION_ACTIONS
from src.engine.reporter import markdown_report
from src.mcp import tools as mcp_tools
from src.web import service as web_service


FORBIDDEN_ACTIONS = {"delete_paragraph", "replace_text", "modify_table_cell_text"}


def _create_report_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("\u6b63\u5f0f\u62a5\u544a")
    document.add_paragraph("\u4e00\u3001\u603b\u4f53\u60c5\u51b5")
    document.add_paragraph("\uff08\u4e00\uff09\u5de5\u4f5c\u8fdb\u5c55")
    document.add_paragraph("1. \u91cd\u70b9\u4e8b\u9879")
    document.add_paragraph("\u8fd9\u662f\u4e00\u6bb5\u6b63\u6587\u5185\u5bb9\u3002")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "\u8868\u683c\u6587\u672c"
    document.save(path)


def _create_report_docx_with_table_body(path: Path) -> None:
    document = Document()
    document.add_paragraph("\u6b63\u5f0f\u62a5\u544a")
    document.add_paragraph("\u4e00\u3001\u603b\u4f53\u60c5\u51b5")
    document.add_paragraph("\uff08\u4e00\uff09\u5de5\u4f5c\u8fdb\u5c55")
    document.add_paragraph("1. \u91cd\u70b9\u4e8b\u9879")
    document.add_paragraph("\u8fd9\u662f\u4e00\u6bb5\u6b63\u6587\u5185\u5bb9\u3002")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "\u8868\u5934"
    table.cell(1, 0).text = "\u8868\u4f53"
    document.save(path)


def test_run_rejects_output_path_equal_to_input_path(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_report_docx(input_docx)

    with pytest.raises(ValueError, match="cannot overwrite input"):
        app.run(str(input_docx), "report", str(input_docx))


def test_run_generates_output_reports_and_only_safe_operations(tmp_path):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    payload = app.run(str(input_docx), "report", str(output_docx))

    assert Path(payload["input"]) == input_docx.resolve()
    assert Path(payload["output"]) == output_docx.resolve()
    assert output_docx.exists()
    assert input_docx.exists()
    assert Path(payload["markdown_report"]).exists()
    assert Path(payload["json_report"]).exists()
    assert Path(payload["llm_analysis"]).exists()
    assert payload["failed_operation_count"] == 0
    assert payload["llm_assistance"]["mode"] == "rule_only_fallback"
    assert payload["llm_assistance"]["operations_source"] == "rule_engine_only"

    operations_payload = json.loads(Path(payload["json_report"]).read_text(encoding="utf-8"))
    actions = {operation["action"] for operation in operations_payload["operations"]}

    assert actions <= SAFE_OPERATION_ACTIONS
    assert not (actions & FORBIDDEN_ACTIONS)


def test_production_entrypoints_use_transactional_workflow():
    for entrypoint in (app.run, web_service.run_format, mcp_tools.apply_docx_template):
        source = inspect.getsource(entrypoint)
        assert "apply_operations_transactional" in source
        assert "apply_operations(" not in source


def test_run_is_idempotent_after_first_formatting_pass(tmp_path):
    input_docx = tmp_path / "input.docx"
    first_output = tmp_path / "first"
    second_output = tmp_path / "second"
    _create_report_docx_with_table_body(input_docx)

    first_payload = app.run(str(input_docx), "report", str(first_output))
    second_payload = app.run(first_payload["output"], "report", str(second_output))

    assert second_payload["issue_count"] == 0
    assert second_payload["operation_count"] == 0
    assert second_payload["failed_operation_count"] == 0


def test_run_failure_does_not_expose_formal_output_and_can_retain_temp(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_dir = tmp_path / "out"
    retained_temp = output_dir / ".retained.tmp.docx"
    _create_report_docx(input_docx)

    def fake_transaction(*args, **kwargs):
        assert kwargs["retain_failed_temp"] is True
        return ExecutionReport(
            status="content_integrity_failed",
            expected_output_path=str(args[1]),
            temp_output_path=str(retained_temp),
            temp_file_retained=True,
            integrity_errors=["paragraph text changed"],
        )

    monkeypatch.setattr(app, "apply_operations_transactional", fake_transaction)

    payload = app.run(
        str(input_docx),
        "report",
        str(output_dir),
        retain_failed_temp=True,
    )

    assert payload["ok"] is False
    assert payload["output"] is None
    assert payload["expected_output"]
    assert payload["temp_file_retained"] is True
    assert payload["temp_output_path"] == str(retained_temp)


def test_run_writes_llm_analysis_json_when_llm_returns_assistance(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class DummyClient:
        def build_assistance(self, document, report):
            return {
                "available": True,
                "mode": "private_llm",
                "status": "ok",
                "document_type": {"result": {"type": "工作报告"}},
                "heading_levels": {"result": {"paragraph_2": "heading_1"}},
                "report_summary": {"result": {"summary": "格式检查完成"}},
                "template_recommendation": {
                    "result": {
                        "recommended_template": "report",
                        "confidence": 0.8,
                        "reason": "文档结构接近报告",
                        "warnings": ["仅供参考"],
                    }
                },
                "operations_source": "rule_engine_only",
                "operations_generated": False,
            }

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    analysis_path = Path(payload["llm_analysis"])
    analysis = json.loads(analysis_path.read_text(encoding="utf-8"))

    assert analysis_path.name == "llm_analysis.json"
    assert analysis["mode"] == "private_llm"
    assert analysis["operations_source"] == "rule_engine_only"
    assert analysis["operations_generated"] is False
    assert analysis["template_recommendation"]["result"]["recommended_template"] == "report"


def test_run_markdown_report_contains_llm_assistance_section(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class DummyClient:
        def build_assistance(self, document, report):
            return {
                "available": True,
                "mode": "private_llm",
                "status": "ok",
                "summary": {"result": {"summary": "仅供人工复核"}},
                "document_type": {"result": {"document_type": "工作报告"}},
                "heading_levels": {"result": {"paragraph_2": "heading_1"}},
                "report_summary": {"result": {"summary": "检查摘要"}},
                "template_recommendation": {
                    "result": {
                        "recommended_template": "memo",
                        "confidence": 0.42,
                        "reason": "部分段落像备忘录",
                        "warnings": ["推荐模板仅供参考"],
                    }
                },
                "operations_source": "rule_engine_only",
                "operations_generated": False,
            }

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    markdown = Path(payload["markdown_report"]).read_text(encoding="utf-8")

    assert "## LLM 辅助分析" in markdown
    assert "LLM 分析仅供参考，不参与格式化操作生成。" in markdown
    assert "- mode: `private_llm`" in markdown
    assert "- operations_source: `rule_engine_only`" in markdown
    assert "- operations_generated: `False`" in markdown
    assert "推荐模板仅供参考，不会覆盖本次实际使用模板。" in markdown
    assert "document_type:" in markdown
    assert "heading_levels:" in markdown
    assert "report_summary:" in markdown
    assert "template_recommendation:" in markdown
    assert "memo" in markdown


def test_run_continues_outputs_when_llm_raises_unhandled_exception(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class BrokenClient:
        def build_assistance(self, document, report):
            raise RuntimeError("private model exploded")

    monkeypatch.setattr(app, "PrivateLLMClient", BrokenClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    analysis = json.loads(Path(payload["llm_analysis"]).read_text(encoding="utf-8"))

    assert output_docx.exists()
    assert Path(payload["markdown_report"]).exists()
    assert Path(payload["json_report"]).exists()
    assert analysis["mode"] == "rule_only_fallback"
    assert analysis["status"] == "fallback"
    assert "RuntimeError: private model exploded" in analysis["error"]
    assert analysis["operations_source"] == "rule_engine_only"
    assert analysis["operations_generated"] is False


def test_operations_json_does_not_include_llm_assistance(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class DummyClient:
        def build_assistance(self, document, report):
            return {
                "available": True,
                "mode": "private_llm",
                "status": "ok",
                "document_type": {"result": {"type": "工作报告"}},
                "operations_source": "rule_engine_only",
                "operations_generated": False,
            }

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    operations_payload = json.loads(Path(payload["json_report"]).read_text(encoding="utf-8"))

    assert "llm_assistance" not in operations_payload
    assert "llm_analysis" not in operations_payload
    assert "document_type" not in operations_payload


def test_llm_template_recommendation_does_not_change_actual_template(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class DummyClient:
        def build_assistance(self, document, report):
            return {
                "available": True,
                "mode": "private_llm",
                "status": "ok",
                "template_recommendation": {
                    "result": {
                        "recommended_template": "not-the-cli-template",
                        "confidence": 0.99,
                        "reason": "LLM advisory recommendation",
                        "warnings": ["must not override CLI template"],
                    }
                },
                "operations_source": "rule_engine_only",
                "operations_generated": False,
            }

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    analysis = json.loads(Path(payload["llm_analysis"]).read_text(encoding="utf-8"))

    assert payload["template"] == "report"
    assert analysis["template_recommendation"]["result"]["recommended_template"] == "not-the-cli-template"


def test_operations_json_ignores_llm_returned_operations(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _create_report_docx(input_docx)

    class DummyClient:
        def build_assistance(self, document, report):
            return {
                "available": True,
                "mode": "private_llm",
                "status": "ok",
                "template_recommendation": {
                    "result": {
                        "recommended_template": "report",
                        "confidence": 0.8,
                        "reason": "advisory only",
                        "warnings": ["LLM operations must be ignored"],
                        "operations": [
                            {
                                "type": "FormatOperation",
                                "action": "apply_paragraph_style",
                                "target_index": 999,
                                "llm_marker": "llm-only-operation",
                            }
                        ],
                    }
                },
                "operations": [
                    {
                        "type": "FormatOperation",
                        "action": "delete_paragraph",
                        "target_index": 999,
                        "llm_marker": "llm-only-operation",
                    }
                ],
                "operations_source": "rule_engine_only",
                "operations_generated": False,
            }

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)

    payload = app.run(str(input_docx), "report", str(output_docx))
    operations_text = Path(payload["json_report"]).read_text(encoding="utf-8")
    operations_payload = json.loads(operations_text)

    assert "llm-only-operation" not in operations_text
    assert "FormatOperation" not in operations_text
    assert "delete_paragraph" not in operations_text
    assert all(operation.get("target_index") != 999 for operation in operations_payload["operations"])
    assert "llm_assistance" not in operations_payload


def test_markdown_report_does_not_depend_on_llm_package():
    source = inspect.getsource(markdown_report)

    assert "src.llm" not in source
    assert "PrivateLLMClient" not in source


def test_run_llm_test_uses_health_check_only(monkeypatch, capsys):
    class DummyConfig:
        enabled = True
        endpoint = "http://private-model:8000/v1"
        model = "local-docx-model"
        timeout = 3
        max_retries = 0

    class DummyClient:
        config = DummyConfig()

        def health_check(self):
            return {
                "available": True,
                "mode": "private_llm",
                "operations_generated": False,
                "operations_source": "rule_engine_only",
            }

    def fail_if_called(*args, **kwargs):
        raise AssertionError("docx pipeline must not run during llm-test")

    monkeypatch.setattr(app, "PrivateLLMClient", DummyClient)
    monkeypatch.setattr(app, "parse_docx", fail_if_called)
    monkeypatch.setattr(app, "check_styles", fail_if_called)
    monkeypatch.setattr(app, "apply_operations_transactional", fail_if_called)

    result = app.run_llm_test()
    output = capsys.readouterr().out

    assert result["available"] is True
    assert "- status: ok" in output
    assert "- operations_source: rule_engine_only" in output
