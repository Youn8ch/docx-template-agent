import json
from pathlib import Path
import re

from docx import Document
from fastapi.testclient import TestClient

from src.llm import assistance
from src.engine.model.operation_model import ExecutionReport
from src.web.app import app
from src.web import service


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Report Title")
    document.add_paragraph("1. First Heading")
    document.add_paragraph("This body text must not change.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text must not change."
    document.save(path)


def _docx_text_snapshot(path: Path) -> tuple[list[str], list[str]]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return paragraphs, cells


def _upload(client: TestClient, route: str, input_docx: Path, data: dict[str, str]):
    with input_docx.open("rb") as file:
        return client.post(
            route,
            data=data,
            files={
                "file": (
                    "input.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )


def test_web_index_lists_upload_form_and_templates():
    client = TestClient(app)

    response = client.get("/")

    assert response.status_code == 200
    assert "docx-template-agent" in response.text
    assert "执行格式检查" in response.text
    assert "执行模板排版" in response.text
    assert 'value="report"' in response.text


def test_web_check_upload_generates_reports(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    client = TestClient(app)

    with input_docx.open("rb") as file:
        response = client.post(
            "/check",
            data={"template_id": "report"},
            files={
                "file": (
                    "input.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    assert "格式检查完成" in response.text
    assert "下载 markdown 检查报告" in response.text
    assert "下载 json 修改明细" in response.text
    assert "下载新 docx" not in response.text

    job_id = re.search(r"/download/([a-f0-9]{32})/markdown", response.text).group(1)
    markdown = client.get(f"/download/{job_id}/markdown")
    detail = client.get(f"/download/{job_id}/json")

    assert markdown.status_code == 200
    assert detail.status_code == 200
    assert markdown.text.startswith("# DOCX")
    assert "operations" in detail.text


def test_web_format_upload_writes_new_docx_without_text_changes(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    before = _docx_text_snapshot(input_docx)
    client = TestClient(app)

    with input_docx.open("rb") as file:
        response = client.post(
            "/format",
            data={"template_id": "report"},
            files={
                "file": (
                    "input.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    assert "模板排版完成" in response.text
    assert "下载新 docx" in response.text

    job_id = re.search(r"/download/([a-f0-9]{32})/docx", response.text).group(1)
    new_docx = client.get(f"/download/{job_id}/docx")
    markdown = client.get(f"/download/{job_id}/markdown")
    detail = client.get(f"/download/{job_id}/json")

    assert new_docx.status_code == 200
    assert markdown.status_code == 200
    assert detail.status_code == 200

    downloaded = tmp_path / "downloaded.docx"
    downloaded.write_bytes(new_docx.content)
    assert _docx_text_snapshot(input_docx) == before
    assert _docx_text_snapshot(downloaded) == before


def test_web_rejects_non_docx_upload(tmp_path):
    input_txt = tmp_path / "input.txt"
    input_txt.write_text("not a docx", encoding="utf-8")
    client = TestClient(app)

    with input_txt.open("rb") as file:
        response = client.post(
            "/format",
            data={"template_id": "report"},
            files={"file": ("input.txt", file, "text/plain")},
        )

    assert response.status_code == 200
    assert "处理失败" in response.text
    assert "only .docx uploads are supported" in response.text


def test_web_check_use_llm_writes_analysis_and_markdown_section(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    client = TestClient(app)

    def fake_build_assistance(document, report):
        return {
            "available": True,
            "mode": "private_llm",
            "status": "ok",
            "document_type": {"result": {"document_type": "report"}},
            "heading_levels": {"result": {"paragraph_1": "heading_1"}},
            "template_recommendation": {"result": {"recommended_template": "other"}},
            "operations_source": "rule_engine_only",
            "operations_generated": False,
        }

    monkeypatch.setattr(service, "build_llm_assistance", fake_build_assistance)

    response = _upload(
        client,
        "/check",
        input_docx,
        {"template_id": "report", "use_llm": "true"},
    )

    assert response.status_code == 200
    job_id = re.search(r"/download/([a-f0-9]{32})/markdown", response.text).group(1)
    markdown = client.get(f"/download/{job_id}/markdown")
    llm = client.get(f"/download/{job_id}/llm_analysis")
    detail = client.get(f"/download/{job_id}/json")

    assert markdown.status_code == 200
    assert llm.status_code == 200
    assert detail.status_code == 200
    assert "LLM 辅助分析" in markdown.text
    analysis = json.loads(llm.text)
    assert analysis["mode"] == "private_llm"
    assert analysis["template_recommendation"]["result"]["recommended_template"] == "other"


def test_web_format_use_llm_writes_all_artifacts_and_keeps_operations_clean(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    before = _docx_text_snapshot(input_docx)
    client = TestClient(app)

    def fake_build_assistance(document, report):
        return {
            "available": True,
            "mode": "private_llm",
            "status": "ok",
            "template_recommendation": {
                "result": {
                    "recommended_template": "not-the-user-template",
                    "operations": [
                        {
                            "type": "FormatOperation",
                            "action": "apply_paragraph_style",
                            "target_index": 999,
                            "llm_marker": "llm-only-operation",
                        }
                    ],
                }
            },
            "operations": [
                {
                    "type": "FormatOperation",
                    "action": "delete_paragraph",
                    "target_index": 999,
                    "llm_marker": "llm-only-operation",
                }
            ],
            "operations_source": "rule_engine_only",
            "operations_generated": False,
        }

    monkeypatch.setattr(service, "build_llm_assistance", fake_build_assistance)

    response = _upload(
        client,
        "/format",
        input_docx,
        {"template_id": "report", "use_llm": "true"},
    )

    assert response.status_code == 200
    job_id = re.search(r"/download/([a-f0-9]{32})/docx", response.text).group(1)
    new_docx = client.get(f"/download/{job_id}/docx")
    markdown = client.get(f"/download/{job_id}/markdown")
    detail = client.get(f"/download/{job_id}/json")
    llm = client.get(f"/download/{job_id}/llm_analysis")

    assert new_docx.status_code == 200
    assert markdown.status_code == 200
    assert detail.status_code == 200
    assert llm.status_code == 200
    assert "LLM 辅助分析" in markdown.text

    downloaded = tmp_path / "downloaded.docx"
    downloaded.write_bytes(new_docx.content)
    assert _docx_text_snapshot(input_docx) == before
    assert _docx_text_snapshot(downloaded) == before

    operations_text = detail.text
    operations_payload = json.loads(operations_text)
    assert "llm_assistance" not in operations_payload
    assert "template_recommendation" not in operations_payload
    assert "llm-only-operation" not in operations_text
    assert "FormatOperation" not in operations_text
    assert all(operation.get("target_index") != 999 for operation in operations_payload["operations"])


def test_web_format_llm_exception_still_generates_formatted_docx(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    before = _docx_text_snapshot(input_docx)
    client = TestClient(app)

    class BrokenClient:
        def build_assistance(self, document, report):
            raise RuntimeError("private model exploded")

    monkeypatch.setattr(assistance, "PrivateLLMClient", BrokenClient)

    response = _upload(
        client,
        "/format",
        input_docx,
        {"template_id": "report", "use_llm": "true"},
    )

    assert response.status_code == 200
    job_id = re.search(r"/download/([a-f0-9]{32})/docx", response.text).group(1)
    new_docx = client.get(f"/download/{job_id}/docx")
    markdown = client.get(f"/download/{job_id}/markdown")
    detail = client.get(f"/download/{job_id}/json")
    llm = client.get(f"/download/{job_id}/llm_analysis")

    assert new_docx.status_code == 200
    assert markdown.status_code == 200
    assert detail.status_code == 200
    assert llm.status_code == 200

    downloaded = tmp_path / "downloaded_after_llm_error.docx"
    downloaded.write_bytes(new_docx.content)
    assert _docx_text_snapshot(downloaded) == before

    analysis = json.loads(llm.text)
    assert analysis["mode"] == "rule_only_fallback"
    assert analysis["status"] == "fallback"
    assert "RuntimeError: private model exploded" in analysis["error"]
    assert analysis["operations_source"] == "rule_engine_only"
    assert analysis["operations_generated"] is False


def test_web_format_failure_has_no_docx_download_and_no_temp_path(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    client = TestClient(app)

    def fake_transaction(*args, **kwargs):
        assert kwargs["retain_failed_temp"] is False
        return ExecutionReport(
            status="content_integrity_failed",
            expected_output_path=str(args[1]),
            temp_output_path=str((tmp_path / ".leaked.tmp.docx").resolve()),
            temp_file_retained=True,
            integrity_errors=["paragraph text changed"],
        )

    monkeypatch.setattr(service, "apply_operations_transactional", fake_transaction)

    response = _upload(client, "/format", input_docx, {"template_id": "report"})

    assert response.status_code == 200
    assert "/download/" in response.text
    assert "/docx" not in response.text
    assert ".leaked.tmp.docx" not in response.text


def test_web_use_llm_false_does_not_call_llm(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    client = TestClient(app)

    def fail_if_called(document, report):
        raise AssertionError("LLM should not be called when use_llm is false")

    monkeypatch.setattr(service, "build_llm_assistance", fail_if_called)

    response = _upload(client, "/check", input_docx, {"template_id": "report"})

    assert response.status_code == 200
    job_id = re.search(r"/download/([a-f0-9]{32})/markdown", response.text).group(1)
    markdown = client.get(f"/download/{job_id}/markdown")
    llm = client.get(f"/download/{job_id}/llm_analysis")

    assert markdown.status_code == 200
    assert "LLM 辅助分析" not in markdown.text
    assert llm.status_code == 404
