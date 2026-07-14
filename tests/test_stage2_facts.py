from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.engine.checker.style_checker import check_styles
from src.engine.parser.docx_parser import parse_docx
from src.engine.parser.facts_extractor import extract_document_facts
from src.engine.parser.structure_detector import detect_structure


def _add_tbl_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_paragraph_outline(paragraph, value: str):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), value)
    p_pr.append(outline)


def _set_style_num_pr(style):
    p_pr = style._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def _template():
    return {
        "template_id": "test",
        "rules": {
            "title": {"font_size": 22},
            "heading_3": {"font_size": 16},
            "body": {"font_size": 12},
            "table": {"font_size": 10.5},
        },
        "safety": {},
    }


def test_body_and_table_paragraphs_use_separate_index_spaces(tmp_path):
    path = tmp_path / "indexes.docx"
    document = Document()
    document.add_paragraph("正文一")
    document.add_paragraph("正文二")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "单元格第一段"
    table.cell(0, 0).add_paragraph("单元格第二段")
    document.save(path)

    facts = extract_document_facts(path, parse_docx(path))
    cell_paragraphs = facts.tables[0].cells[0].paragraphs

    assert [item.paragraph_index for item in facts.body_paragraphs] == [1, 2]
    assert all(item.container_type == "body" for item in facts.body_paragraphs)
    assert [item.paragraph_index for item in cell_paragraphs] == [None, None]
    assert [item.cell_paragraph_index for item in cell_paragraphs] == [1, 2]
    assert cell_paragraphs[0].table_index == 1
    assert cell_paragraphs[0].row_index == 1
    assert cell_paragraphs[0].cell_index == 1


def test_mixed_run_format_summary_distinguishes_local_bold(tmp_path):
    path = tmp_path / "runs.docx"
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Alpha")
    first.font.name = "Arial"
    first.font.size = Pt(12)
    second = paragraph.add_run("Beta")
    second.font.name = "SimSun"
    second.font.size = Pt(16)
    second.bold = True
    document.save(path)

    facts = extract_document_facts(path)
    summary = facts.body_paragraphs[0].run_format_summary

    assert summary.run_count == 2
    assert summary.non_empty_run_count == 2
    assert set(summary.font_names) == {"Arial", "SimSun"}
    assert summary.mixed_font is True
    assert set(summary.font_sizes) == {12.0, 16.0}
    assert summary.mixed_font_size is True
    assert summary.any_bold is True
    assert summary.all_bold is False


def test_outline_and_heading_style_hard_facts_do_not_change_default_role(tmp_path):
    path = tmp_path / "heading-style.docx"
    document = Document()
    style = document.styles.add_style("自定义标题", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Heading 1"]
    paragraph = document.add_paragraph("不是编号但继承标题样式", style="自定义标题")
    _set_paragraph_outline(paragraph, "0")
    document.save(path)

    parsed_default = detect_structure(parse_docx(path))
    parsed_with_facts = parse_docx(path, include_facts=True)
    fact = parsed_with_facts.facts.body_paragraphs[0]

    assert parsed_default.paragraphs[0].role == "title"
    assert parsed_with_facts.paragraphs[0].role == "unknown"
    assert fact.word_features.heading_level == 1
    assert fact.word_features.inherited_trusted_heading_style is True
    assert fact.hard_role_result.role_type == "outline_heading"


def test_toc_static_entries_form_continuous_region(tmp_path):
    path = tmp_path / "toc.docx"
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("第一章 总则....................1")
    document.add_paragraph("第二章 细则\t2")
    document.add_paragraph("正文开始")
    document.save(path)

    facts = extract_document_facts(path)
    region = facts.regions[0]

    assert region.type == "toc"
    assert region.start_paragraph_index == 1
    assert region.end_paragraph_index == 3
    assert "toc_heading_text" in region.evidence
    assert facts.body_paragraphs[0].region_flags.toc_heading_candidate is True
    assert facts.body_paragraphs[1].region_flags.toc_entry_candidate is True
    assert facts.body_paragraphs[2].hard_role_result.role_type == "toc_entry"
    assert facts.body_paragraphs[3].region_flags.toc_region is False


def test_toc_field_and_style_are_entry_candidates(tmp_path):
    path = tmp_path / "toc-field.docx"
    document = Document()
    document.add_paragraph("目录")
    field = document.add_paragraph()
    run = field.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1-3" \\h'
    run._r.append(instr)
    try:
        document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    styled = document.add_paragraph("Styled entry 3")
    styled.style = "TOC 1"
    document.save(path)

    facts = extract_document_facts(path)

    assert facts.body_paragraphs[1].region_flags.toc_entry_candidate is True
    assert "toc_field" in facts.body_paragraphs[1].context["toc_entry_evidence"]
    assert facts.body_paragraphs[2].region_flags.toc_entry_candidate is True
    assert "toc_entry_style" in facts.body_paragraphs[2].context["toc_entry_evidence"]


def test_table_header_candidate_carries_evidence(tmp_path):
    path = tmp_path / "table-header.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    _add_tbl_header(table.rows[0])
    for cell in table.rows[0].cells:
        _shade_cell(cell, "D9EAF7")
        run = cell.paragraphs[0].add_run("Head")
        run.bold = True
    table.cell(1, 0).text = "Longer body text"
    table.cell(1, 1).text = "Another longer body text"
    document.save(path)

    facts = extract_document_facts(path)
    candidate = facts.tables[0].header_candidate

    assert candidate.is_candidate is True
    assert candidate.score > 0.5
    assert "first_row" in candidate.evidence
    assert "tblHeader" in candidate.evidence
    assert "all_cells_bold" in candidate.evidence
    assert "distinct_shading" in candidate.evidence
    assert "shorter_text_than_body_rows" in candidate.evidence
    assert facts.tables[0].cells[0].paragraphs[0].region_flags.table_header_candidate is True


def test_facts_serialization_is_stable(tmp_path):
    path = tmp_path / "stable.docx"
    document = Document()
    document.add_paragraph("Hello")
    document.save(path)

    first = extract_document_facts(path).model_dump(mode="json")
    second = extract_document_facts(path).model_dump(mode="json")

    assert first == second


def test_parse_docx_default_behavior_and_operations_stay_unchanged(tmp_path):
    path = tmp_path / "compat.docx"
    document = Document()
    paragraph = document.add_paragraph("1. Heading")
    paragraph.add_run().font.size = Pt(10)
    document.save(path)

    default_model = detect_structure(parse_docx(path))
    facts_model = parse_docx(path, include_facts=True)
    report_default = check_styles(default_model, _template())
    report_after_facts = check_styles(default_model, _template())

    assert parse_docx(path).facts is None
    assert facts_model.facts is not None
    assert report_default.model_dump(mode="json") == report_after_facts.model_dump(mode="json")


def test_trusted_heading_matching_is_exact(tmp_path):
    path = tmp_path / "heading-boundaries.docx"
    document = Document()
    for name in ["标题 10", "标题 1 custom", "Heading 1 custom"]:
        document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        document.add_paragraph(name, style=name)
    document.save(path)

    facts = extract_document_facts(path)

    assert [item.word_features.trusted_builtin_heading for item in facts.body_paragraphs] == [False, False, False]
    assert [item.hard_role_result for item in facts.body_paragraphs] == [None, None, None]


def test_heading_outline_conflict_downgrades_inherited_heading(tmp_path):
    path = tmp_path / "heading-conflict.docx"
    document = Document()
    inherited = document.styles.add_style("Inherited Heading", WD_STYLE_TYPE.PARAGRAPH)
    inherited.base_style = document.styles["Heading 1"]
    paragraph = document.add_paragraph("Inherited but body outline", style="Inherited Heading")
    _set_paragraph_outline(paragraph, "9")
    document.save(path)

    fact = extract_document_facts(path).body_paragraphs[0]

    assert fact.word_features.inherited_heading_evidence is True
    assert fact.word_features.outline_conflict is True
    assert fact.hard_role_result is None
    assert any(hint.hint_type == "outline_conflict" for hint in fact.rule_hints)


def test_inherited_heading_without_explicit_outline_is_evidence_only(tmp_path):
    path = tmp_path / "inherited-heading-evidence-only.docx"
    document = Document()
    inherited = document.styles.add_style("Evidence Only Heading", WD_STYLE_TYPE.PARAGRAPH)
    inherited.base_style = document.styles["Heading 1"]
    document.add_paragraph("Inherited heading evidence only", style="Evidence Only Heading")
    document.save(path)

    fact = extract_document_facts(path).body_paragraphs[0]

    assert fact.word_features.inherited_heading_evidence is True
    assert fact.word_features.explicit_outline_level is False
    assert fact.word_features.inherited_outline_level is True
    assert fact.hard_role_result is None
    assert any(hint.hint_type == "inherited_heading_evidence" for hint in fact.rule_hints)


def test_inherited_heading_with_explicit_outline_uses_explicit_source(tmp_path):
    path = tmp_path / "inherited-heading-explicit-outline.docx"
    document = Document()
    inherited = document.styles.add_style("Explicit Outline Heading", WD_STYLE_TYPE.PARAGRAPH)
    inherited.base_style = document.styles["Heading 1"]
    outline_2 = document.add_paragraph("Explicit level 2", style="Explicit Outline Heading")
    _set_paragraph_outline(outline_2, "1")
    outline_3 = document.add_paragraph("Explicit level 3", style="Explicit Outline Heading")
    _set_paragraph_outline(outline_3, "2")
    document.save(path)

    first, second = extract_document_facts(path).body_paragraphs

    assert first.word_features.heading_level == 2
    assert first.hard_role_result.role_type == "outline_heading"
    assert first.hard_role_result.evidence == ["explicit_outline_level"]
    assert second.word_features.heading_level == 3
    assert second.hard_role_result.role_type == "outline_heading"
    assert second.hard_role_result.evidence == ["explicit_outline_level"]


def test_builtin_heading_with_explicit_outline_conflict_records_evidence(tmp_path):
    path = tmp_path / "builtin-heading-outline-conflict.docx"
    document = Document()
    paragraph = document.add_paragraph("Builtin style with explicit level 2", style="Heading 1")
    _set_paragraph_outline(paragraph, "1")
    document.save(path)

    fact = extract_document_facts(path).body_paragraphs[0]

    assert fact.word_features.trusted_builtin_heading is True
    assert fact.word_features.heading_level == 2
    assert fact.hard_role_result.role_type == "outline_heading"
    assert fact.hard_role_result.evidence == ["explicit_outline_level"]
    assert fact.context["heading_conflicts"] == [
        {
            "type": "heading_style_outline_conflict",
            "style_heading_level": 1,
            "explicit_outline_level": 2,
        }
    ]
    assert any(hint.hint_type == "heading_style_outline_conflict" for hint in fact.rule_hints)


def test_heading_outline_rules_for_builtin_and_inherited(tmp_path):
    path = tmp_path / "heading-rules.docx"
    document = Document()
    builtin = document.add_paragraph("Built in", style="Heading 1")
    _set_paragraph_outline(builtin, "0")
    inherited_style = document.styles.add_style("Inherited Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    inherited_style.base_style = document.styles["Heading 1"]
    inherited_valid = document.add_paragraph("Inherited valid outline", style="Inherited Heading 2")
    _set_paragraph_outline(inherited_valid, "1")
    document.save(path)

    facts = extract_document_facts(path)

    assert facts.body_paragraphs[0].word_features.trusted_builtin_heading is True
    assert facts.body_paragraphs[0].hard_role_result.role_type == "outline_heading"
    assert facts.body_paragraphs[1].word_features.inherited_heading_evidence is True
    assert facts.body_paragraphs[1].word_features.explicit_outline_level is True
    assert facts.body_paragraphs[1].hard_role_result.role_type == "outline_heading"


def test_toc_conservative_candidate_and_confirmed_regions(tmp_path):
    single_path = tmp_path / "toc-single.docx"
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("第一章 总则....................1")
    document.add_paragraph("正文")
    document.save(single_path)

    single = extract_document_facts(single_path)
    assert single.regions == []
    assert single.body_paragraphs[0].region_flags.toc_region_candidate is True
    assert single.body_paragraphs[1].region_flags.toc_region_candidate is True
    assert single.body_paragraphs[1].region_flags.confirmed_toc_region is False

    confirmed_path = tmp_path / "toc-confirmed.docx"
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("第一章 总则 1")
    document.add_paragraph("第二章 细则 2")
    document.add_paragraph("")
    document.add_paragraph("正文 2026")
    document.save(confirmed_path)

    confirmed = extract_document_facts(confirmed_path)
    assert len(confirmed.regions) == 1
    assert confirmed.regions[0].start_paragraph_index == 1
    assert confirmed.regions[0].end_paragraph_index == 3
    assert confirmed.regions[0].confidence == 0.9
    assert confirmed.body_paragraphs[4].region_flags.toc_region is False


def test_toc_style_and_field_confirmation(tmp_path):
    style_path = tmp_path / "toc-style.docx"
    document = Document()
    try:
        document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    document.add_paragraph("Entry A 1", style="TOC 1")
    document.add_paragraph("Entry B 2", style="TOC 1")
    document.save(style_path)
    style_facts = extract_document_facts(style_path)

    assert len(style_facts.regions) == 1
    assert style_facts.regions[0].evidence == ["toc_entry_style"]

    field_path = tmp_path / "toc-field-only.docx"
    document = Document()
    field = document.add_paragraph()
    run = field.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1-3" \\h'
    run._r.append(instr)
    document.save(field_path)
    field_facts = extract_document_facts(field_path)

    assert len(field_facts.regions) == 1
    assert field_facts.regions[0].confidence == 0.95
    assert "toc_field" in field_facts.regions[0].evidence
    assert field_facts.body_paragraphs[0].hard_role_result.role_type == "toc_entry"
    assert "toc_field" in field_facts.body_paragraphs[0].hard_role_result.evidence


def test_confirmed_toc_field_overrides_empty_hard_role(tmp_path):
    path = tmp_path / "toc-field-empty.docx"
    document = Document()
    field = document.add_paragraph()
    run = field.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1-3" \\h'
    run._r.append(instr)
    document.add_paragraph("")
    document.save(path)

    facts = extract_document_facts(path)

    assert facts.body_paragraphs[0].region_flags.confirmed_toc_region is True
    assert facts.body_paragraphs[0].hard_role_result.role_type == "toc_entry"
    assert "toc_field" in facts.body_paragraphs[0].hard_role_result.evidence
    assert facts.body_paragraphs[1].hard_role_result.role_type == "empty"


def test_toc_field_with_following_entries_keeps_consistent_region_roles(tmp_path):
    path = tmp_path / "toc-field-with-entries.docx"
    document = Document()
    document.add_paragraph("目录")
    field = document.add_paragraph()
    run = field.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1-3" \\h'
    run._r.append(instr)
    document.add_paragraph("第一章 总则................1")
    document.save(path)

    facts = extract_document_facts(path)

    field_fact = facts.body_paragraphs[1]
    assert field_fact.region_flags.confirmed_toc_region is True
    assert field_fact.hard_role_result.role_type == "toc_entry"
    assert "toc_field" in field_fact.hard_role_result.evidence
    assert facts.body_paragraphs[2].region_flags.toc_entry_candidate is True


def test_merged_and_nested_table_metadata(tmp_path):
    merged_path = tmp_path / "merged.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "merged"
    table.cell(1, 0).merge(table.cell(1, 1))
    document.save(merged_path)

    parsed = parse_docx(merged_path, include_facts=True)
    table_facts = parsed.facts.tables[0]

    assert table_facts.exposure_model == "physical_w_tc"
    assert table_facts.legacy_cell_index_compatible is False
    assert table_facts.contains_merged_cells is True
    assert table_facts.cells[0].physical_cell_index == 1
    assert table_facts.cells[0].logical_grid_start == 1
    assert table_facts.cells[0].grid_span == 2
    assert table_facts.cells[0].is_merged is True
    assert len(parsed.tables[0].cells) != len(table_facts.cells)

    nested_path = tmp_path / "nested.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "outer"
    table.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "inner"
    document.save(nested_path)
    nested = extract_document_facts(nested_path).tables[0]

    assert nested.contains_nested_tables is True
    assert nested.nested_tables_extracted is False
    assert nested.cells[0].contains_nested_table is True
    assert nested.cells[0].nested_table_count == 1
    assert "nested_table" in nested.cells[0].unsupported_features


def test_vmerge_continuation_points_to_restart_anchor(tmp_path):
    path = tmp_path / "vmerge-anchor.docx"
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "vertical"
    document.save(path)

    table_facts = extract_document_facts(path).tables[0]
    restart, continuation = table_facts.cells

    assert restart.merge_kind == "v_merge_restart"
    assert restart.merged_anchor == "r1c1"
    assert restart.anchor_resolved is True
    assert continuation.merge_kind == "v_merge_continue"
    assert continuation.merged_anchor == "r1c1"
    assert continuation.anchor_resolved is True


def test_mixed_gridspan_vmerge_continuation_does_not_self_anchor(tmp_path):
    path = tmp_path / "mixed-merge-anchor.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 1))
    table.cell(0, 0).text = "block"
    document.save(path)

    cells = extract_document_facts(path).tables[0].cells
    restart = cells[0]
    continuation = [cell for cell in cells if cell.v_merge == "continue"][0]

    assert restart.merge_kind == "mixed"
    assert restart.merged_anchor == "r1c1"
    assert restart.anchor_resolved is True
    assert continuation.merge_kind == "mixed"
    assert continuation.merged_anchor == "r1c1"
    assert continuation.merged_anchor != f"r{continuation.row_index}c{continuation.logical_grid_start}"
    assert continuation.anchor_resolved is True


def test_unresolved_vmerge_anchor_is_explicit(tmp_path):
    path = tmp_path / "unresolved-vmerge.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
    vmerge = OxmlElement("w:vMerge")
    tc_pr.append(vmerge)
    document.save(path)

    cell = extract_document_facts(path).tables[0].cells[0]

    assert cell.merge_kind == "v_merge_continue"
    assert cell.merged_anchor is None
    assert cell.anchor_resolved is False
    assert "vmerge_anchor_unresolved" in cell.unsupported_features


def test_run_summary_dedup_weighting_channels_and_whitespace(tmp_path):
    path = tmp_path / "run-summary.docx"
    document = Document()
    paragraph = document.add_paragraph()
    long_run = paragraph.add_run("中" * 100)
    long_run.font.name = "Arial"
    long_run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun")
    short_run = paragraph.add_run("x")
    short_run.font.name = "Calibri"
    short_run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimHei")
    paragraph.add_run(" \t\n")
    repeat = paragraph.add_run("tail")
    repeat.font.name = "Arial"
    repeat._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun")
    document.save(path)

    summary = extract_document_facts(path).body_paragraphs[0].run_format_summary

    assert summary.non_empty_run_count == 3
    assert summary.whitespace_run_count == 1
    assert summary.font_names == ["SimSun", "SimHei"]
    assert summary.east_asia_font_names == ["SimSun", "SimHei"]
    assert summary.ascii_font_names == ["Arial", "Calibri"]
    assert summary.dominant_font_name == "SimSun"
    assert summary.dominant_east_asia_font == "SimSun"
    assert summary.dominant_ascii_font == "Arial"


def test_numbering_style_scope_is_explicit(tmp_path):
    path = tmp_path / "style-numbering.docx"
    document = Document()
    style = document.styles.add_style("NumberedLike", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_num_pr(style)
    document.add_paragraph("Styled numbering", style="NumberedLike")
    document.save(path)

    numbering = extract_document_facts(path).body_paragraphs[0].word_features.numbering

    assert numbering.direct_numbering_present is False
    assert numbering.style_numbering_supported is False
    assert numbering.numbering_resolution_scope == "paragraph_direct_only"
    assert numbering.unsupported_reason == "style-derived numbering is not resolved in stage 2"


def test_preview_limits_and_internal_facts_marker(tmp_path):
    path = tmp_path / "preview.docx"
    document = Document()
    document.add_paragraph("x" * 200)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "y" * 200
    document.save(path)

    facts = extract_document_facts(path)

    assert facts.internal_only is True
    assert facts.llm_payload_safe is False
    assert facts.snapshot_required is True
    assert facts.preview_max_chars == 120
    assert len(facts.body_paragraphs[0].text_preview) == 120
    assert len(facts.tables[0].cells[0].text_preview) == 120
    assert len(facts.tables[0].cells[0].paragraphs[0].text_preview) == 120
