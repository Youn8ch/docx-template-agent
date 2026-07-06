from pathlib import Path
from tempfile import TemporaryDirectory
import json
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from src.engine.formatter.apply_template import apply_operations
from src.engine.formatter.font_utils import first_line_indent_chars_value, is_valid_font_size, normalize_font_name
from src.engine.model.operation_model import (
    FormatOperation,
    is_safe_operation,
)
from src.engine.model.document_model import DocumentModel, ParagraphInfo, TableCellInfo, TableInfo
from src.engine.parser.docx_parser import parse_docx
from src.engine.parser.structure_detector import detect_structure
from src.engine.reporter.json_report import write_json_report
from src.engine.reporter.markdown_report import write_markdown_report
from src.engine.checker.style_checker import check_styles
from src.engine.template.template_loader import list_templates, load_template


class SkeletonTest(unittest.TestCase):
    def test_template_loads(self):
        template = load_template("report")
        self.assertEqual(template["template_id"], "report")
        self.assertIn("heading_1", template["rules"])
        self.assertFalse(template["safety"]["allow_text_change"])

    def test_list_templates_returns_metadata(self):
        templates = list_templates()
        self.assertIn(
            {
                "template_id": "report",
                "template_name": templates[0]["template_name"],
                "description": templates[0]["description"],
            },
            templates,
        )

    def test_safe_operation_allowed(self):
        operation = FormatOperation(
            operation_id="op-0001",
            action="apply_paragraph_style",
            target_type="paragraph",
            target_indices=[0],
        )
        self.assertTrue(is_safe_operation(operation))

    def test_dangerous_operation_rejected(self):
        with self.assertRaises(ValueError):
            FormatOperation(
                operation_id="op-0002",
                action="replace_text",
                target_type="paragraph",
                target_indices=[0],
            )

    def test_font_utils_are_conservative(self):
        self.assertEqual(normalize_font_name(" SimSun "), "SimSun")
        self.assertIsNone(normalize_font_name(" "))
        self.assertTrue(is_valid_font_size(12))
        self.assertFalse(is_valid_font_size(500))
        self.assertEqual(first_line_indent_chars_value({"first_line_indent_chars": 2}), 200)
        self.assertEqual(first_line_indent_chars_value({"first_line_indent_chars": 0}), 0)

    def test_parse_docx_extracts_structure_and_formats(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "sample.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph._p.get_or_add_pPr().ind.set(qn("w:firstLineChars"), "200")
            run = paragraph.add_run("Hello")
            run.font.name = "SimSun"
            run.font.size = Pt(12)
            run.font.bold = True

            document.add_paragraph()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            document.save(path)

            model = parse_docx(path)

        self.assertEqual(model.filepath, path)
        self.assertEqual(model.paragraph_count, 2)
        self.assertEqual(model.table_count, 1)
        self.assertEqual(model.paragraphs[0].index, 1)
        self.assertEqual(model.paragraphs[0].text, "Hello")
        self.assertEqual(model.paragraphs[0].style_name, "Normal")
        self.assertEqual(model.paragraphs[0].alignment, "center")
        self.assertEqual(model.paragraphs[0].line_spacing, 1.5)
        self.assertEqual(model.paragraphs[0].space_before, 6)
        self.assertEqual(model.paragraphs[0].space_after, 12)
        self.assertEqual(model.paragraphs[0].first_line_indent, 24)
        self.assertEqual(model.paragraphs[0].first_line_indent_chars, 2)
        self.assertEqual(model.paragraphs[0].runs[0].font_name, "SimSun")
        self.assertEqual(model.paragraphs[0].runs[0].font_size, 12)
        self.assertTrue(model.paragraphs[0].runs[0].bold)
        self.assertEqual(model.paragraphs[1].index, 2)
        self.assertEqual(model.paragraphs[1].text, "")
        self.assertEqual(model.paragraphs[1].runs, [])
        self.assertEqual(model.tables[0].index, 1)
        self.assertEqual(model.tables[0].rows, 1)
        self.assertEqual(model.tables[0].cols, 2)
        self.assertEqual(model.tables[0].cells[0].text, "A")
        self.assertIn("Normal", model.styles)

    def test_detect_structure_from_paragraph_text(self):
        document = DocumentModel(
            filepath=Path("sample.docx"),
            paragraphs=[
                ParagraphInfo(index=1, text="   "),
                ParagraphInfo(index=2, text="正式报告"),
                ParagraphInfo(index=3, text="一、总体情况"),
                ParagraphInfo(index=4, text="（一）工作进展"),
                ParagraphInfo(index=5, text="1. 重点事项"),
                ParagraphInfo(index=6, text="1、另一事项"),
                ParagraphInfo(index=7, text="1.1 分项说明"),
                ParagraphInfo(index=8, text="这是一段正文内容。"),
            ],
        )

        detected = detect_structure(document)

        self.assertEqual(
            [paragraph.role for paragraph in detected.paragraphs],
            [
                "empty",
                "title",
                "heading_1",
                "heading_2",
                "heading_3",
                "heading_3",
                "heading_3",
                "body",
            ],
        )
        self.assertEqual(
            [paragraph.text for paragraph in detected.paragraphs],
            [paragraph.text for paragraph in document.paragraphs],
        )

    def test_check_styles_generates_report_and_operations(self):
        template = {
            "template_id": "test",
            "rules": {
                "body": {
                    "font_name": "SimSun",
                    "font_size": 12,
                    "bold": False,
                    "alignment": "justify",
                    "line_spacing": 1.5,
                    "space_before": 0,
                    "space_after": 0,
                    "first_line_indent_chars": 2,
                },
                "table": {
                    "font_name": "SimSun",
                    "font_size": 10.5,
                    "bold": False,
                    "alignment": "center",
                },
                "table_header": {
                    "font_name": "SimHei",
                    "font_size": 10.5,
                    "bold": True,
                    "alignment": "center",
                },
            },
            "safety": {},
        }
        document = DocumentModel(
            filepath=Path("sample.docx"),
            paragraphs=[
                ParagraphInfo(index=1, text="", role="empty"),
                ParagraphInfo(
                    index=2,
                    text="body",
                    role="body",
                    alignment="left",
                    font_names=["Arial"],
                    font_sizes=[10],
                    bold_values=[True],
                    line_spacing=1.0,
                    space_before=3,
                    space_after=6,
                    first_line_indent=0,
                ),
            ],
            tables=[
                TableInfo(
                    index=1,
                    rows=2,
                    cols=1,
                    cells=[
                        TableCellInfo(
                            row_index=1,
                            col_index=1,
                            text="header",
                            alignment="left",
                            font_names=["Arial"],
                            font_sizes=[9],
                            bold_values=[False],
                        ),
                        TableCellInfo(
                            row_index=2,
                            col_index=1,
                            text="body",
                            alignment="left",
                            font_names=["Arial"],
                            font_sizes=[9],
                            bold_values=[True],
                        ),
                    ],
                )
            ],
        )

        report = check_styles(document, template)

        self.assertEqual(report.template_id, "test")
        self.assertGreater(report.issue_count, 0)
        self.assertGreaterEqual(report.operation_count, 3)
        self.assertNotIn(1, [operation.target_indices[0] for operation in report.operations if operation.target_type == "paragraph"])
        self.assertTrue(all(is_safe_operation(operation) for operation in report.operations))
        self.assertIn("apply_paragraph_style", [operation.action for operation in report.operations])
        self.assertIn("apply_table_style", [operation.action for operation in report.operations])
        self.assertIn("apply_table_header_style", [operation.action for operation in report.operations])

    def test_apply_operations_writes_formatted_copy(self):
        with TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(30)
            paragraph.add_run("Paragraph")
            heading = document.add_paragraph()
            heading.paragraph_format.first_line_indent = Pt(24)
            heading.paragraph_format.left_indent = Pt(18)
            heading.add_run("Heading")
            table = document.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Header"
            table.cell(1, 0).text = "Body"
            table.cell(1, 0).paragraphs[0].paragraph_format.first_line_indent = Pt(24)
            document.save(source)

            operations = [
                FormatOperation(
                    operation_id="op-0001",
                    action="apply_paragraph_style",
                    target_type="paragraph",
                    target_indices=[1],
                    role="body",
                    properties={
                        "font_name": "SimSun",
                        "font_size": 12,
                        "bold": True,
                        "alignment": "center",
                        "line_spacing": 1.5,
                        "space_before": 6,
                        "space_after": 9,
                        "first_line_indent_chars": 2,
                    },
                ),
                FormatOperation(
                    operation_id="op-0002",
                    action="apply_paragraph_style",
                    target_type="paragraph",
                    target_indices=[2],
                    role="title",
                    properties={
                        "font_name": "SimHei",
                        "font_size": 22,
                        "bold": True,
                        "alignment": "center",
                        "first_line_indent_chars": 0,
                    },
                ),
                FormatOperation(
                    operation_id="op-0003",
                    action="apply_table_style",
                    target_type="table",
                    target_indices=[1],
                    properties={
                        "font_name": "SimSun",
                        "font_size": 10.5,
                        "alignment": "center",
                    },
                ),
                FormatOperation(
                    operation_id="op-0004",
                    action="apply_table_header_style",
                    target_type="table_header",
                    target_indices=[1],
                    properties={
                        "font_name": "SimHei",
                        "font_size": 10.5,
                        "bold": True,
                        "alignment": "center",
                    },
                ),
            ]

            results = apply_operations(source, output, operations)
            formatted = Document(str(output))

        self.assertTrue(all(result["status"] == "success" for result in results))
        run = formatted.paragraphs[0].runs[0]
        self.assertEqual(run.font.name, "SimSun")
        self.assertEqual(run._element.rPr.rFonts.get(qn("w:eastAsia")), "SimSun")
        self.assertEqual(run.font.size.pt, 12)
        self.assertTrue(run.font.bold)
        self.assertEqual(formatted.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(formatted.paragraphs[0].paragraph_format.line_spacing, 1.5)
        self.assertEqual(formatted.paragraphs[0].paragraph_format.space_before.pt, 6)
        self.assertEqual(formatted.paragraphs[0].paragraph_format.space_after.pt, 9)
        body_indent = formatted.paragraphs[0]._p.pPr.ind
        self.assertEqual(body_indent.get(qn("w:firstLineChars")), "200")
        self.assertIsNone(body_indent.get(qn("w:firstLine")))
        self.assertIsNone(body_indent.get(qn("w:hanging")))
        self.assertEqual(body_indent.get(qn("w:left")), "0")

        heading_indent = formatted.paragraphs[1]._p.pPr.ind
        self.assertEqual(heading_indent.get(qn("w:firstLineChars")), "0")
        self.assertIsNone(heading_indent.get(qn("w:firstLine")))
        self.assertIsNone(heading_indent.get(qn("w:hanging")))
        self.assertIsNone(heading_indent.get(qn("w:hangingChars")))
        self.assertEqual(heading_indent.get(qn("w:left")), "0")

        header_run = formatted.tables[0].cell(0, 0).paragraphs[0].runs[0]
        body_run = formatted.tables[0].cell(1, 0).paragraphs[0].runs[0]
        self.assertEqual(header_run.font.name, "SimHei")
        self.assertEqual(header_run._element.rPr.rFonts.get(qn("w:eastAsia")), "SimHei")
        self.assertTrue(header_run.font.bold)
        self.assertEqual(body_run.font.name, "SimSun")
        self.assertEqual(body_run._element.rPr.rFonts.get(qn("w:eastAsia")), "SimSun")
        self.assertEqual(formatted.tables[0].cell(1, 0).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        table_body_indent = formatted.tables[0].cell(1, 0).paragraphs[0]._p.pPr.ind
        self.assertIsNone(table_body_indent.get(qn("w:firstLineChars")))
        self.assertIsNone(table_body_indent.get(qn("w:firstLine")))
        self.assertIsNone(table_body_indent.get(qn("w:hanging")))
        self.assertEqual([paragraph.text for paragraph in formatted.paragraphs], ["Paragraph", "Heading"])
        self.assertEqual(formatted.tables[0].cell(0, 0).text, "Header")
        self.assertEqual(formatted.tables[0].cell(1, 0).text, "Body")

    def test_reporters_generate_paths_from_output_docx(self):
        template = {
            "template_id": "test",
            "rules": {"body": {"font_name": "SimSun"}},
            "safety": {},
        }
        document = DocumentModel(
            filepath=Path("source.docx"),
            paragraphs=[
                ParagraphInfo(index=1, text="body", role="body", font_names=["Arial"]),
            ],
        )
        report = check_styles(document, template)

        with TemporaryDirectory() as temp_dir:
            output_docx = Path(temp_dir) / "source_formatted.docx"
            markdown_path = write_markdown_report(
                report,
                output_docx,
                "source.docx",
                "test",
                document.paragraph_count,
                document.table_count,
            )
            json_path = write_json_report(report, output_docx)
            markdown = markdown_path.read_text(encoding="utf-8")
            payload = json.loads(json_path.read_text(encoding="utf-8"))

        self.assertEqual(markdown_path.name, "source_formatted_check_report.md")
        self.assertEqual(json_path.name, "source_formatted_operations.json")
        self.assertIn("原始文件", markdown)
        self.assertIn("问题分类统计", markdown)
        self.assertIn("问题明细", markdown)
        self.assertIn("人工复核建议", markdown)
        self.assertEqual(payload["template_id"], "test")
        self.assertEqual(payload["issue_count"], report.issue_count)
        self.assertEqual(payload["operation_count"], report.operation_count)
        self.assertEqual(len(payload["issues"]), report.issue_count)
        self.assertEqual(len(payload["operations"]), report.operation_count)


if __name__ == "__main__":
    unittest.main()
