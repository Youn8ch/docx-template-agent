import importlib.util
import json
from pathlib import Path

from docx import Document
import pytest

from src.mcp import tools


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Report Title")
    document.add_paragraph("1. First Heading")
    document.add_paragraph("This body text must not change.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text must not change."
    document.save(path)


def _docx_text_snapshot(path: Path) -> tuple[list[str], list[str]]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return paragraphs, cells


def _cleanup_result_files(result: dict) -> None:
    for key in ("output_docx_path", "markdown_report_path", "json_detail_path", "llm_analysis_path"):
        value = result.get(key)
        if value:
            Path(value).unlink(missing_ok=True)


def _fake_llm_assistance() -> dict:
    return {
        "available": True,
        "mode": "private_llm",
        "status": "ok",
        "document_type": {"result": {"type": "report"}},
        "heading_levels": {"result": [{"index": 0, "level": 1}]},
        "report_summary": {"result": {"summary": "mcp llm summary"}},
        "template_recommendation": {
            "result": {
                "recommended_template": "other-template",
                "operations": [{"llm_marker": "llm-only-operation"}],
            }
        },
        "operations": [{"llm_marker": "llm-top-level-operation"}],
        "operations_generated": False,
        "operations_source": "rule_engine_only",
    }


def test_mcp_server_registers_expected_tools():
    pytest.importorskip("mcp")
    project_root = Path(__file__).resolve().parents[1]
    spec = importlib.util.spec_from_file_location(
        "project_mcp_server",
        project_root / "mcp_server.py",
    )
    assert spec is not None
    assert spec.loader is not None
    mcp_server = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mcp_server)

    assert set(mcp_server.mcp._tool_manager._tools) == {
        "list_templates",
        "analyze_docx",
        "check_docx_style",
        "apply_docx_template",
        "generate_review_report",
    }


def test_mcp_tools_return_json_for_list_analyze_and_check(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    templates = tools.list_templates({})
    analysis = tools.analyze_docx({"input_path": str(input_docx), "include_tables": False})
    check = tools.check_docx_style(
        {
            "input_path": str(input_docx),
            "template": "report",
            "include_issues": False,
            "include_operations": False,
        }
    )

    assert templates["ok"] is True
    assert templates["count"] >= 1
    assert analysis["ok"] is True
    assert analysis["paragraph_count"] == 3
    assert analysis["table_count"] == 1
    assert check["ok"] is True
    assert check["issue_count"] >= 0
    assert check["operation_count"] >= 0
    assert check["issues"] == []
    assert check["operations"] == []


def test_apply_docx_template_writes_allowed_outputs_without_text_changes(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    before = _docx_text_snapshot(input_docx)

    result = tools.apply_docx_template({"input_path": str(input_docx), "template": "report"})

    try:
        assert result["ok"] is True
        output_docx = Path(result["output_docx_path"])
        markdown_report = Path(result["markdown_report_path"])
        json_detail = Path(result["json_detail_path"])

        assert output_docx.exists()
        assert markdown_report.exists()
        assert json_detail.exists()
        assert output_docx != input_docx.resolve()
        assert "samples\\output" in str(output_docx) or "samples/output" in str(output_docx)
        assert result["issue_count"] >= 0
        assert result["operation_count"] >= 0
        assert "llm_analysis_path" not in result
        assert _docx_text_snapshot(input_docx) == before
        assert _docx_text_snapshot(output_docx) == before
    finally:
        _cleanup_result_files(result)


def test_generate_review_report_does_not_create_docx(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    result = tools.generate_review_report({"input_path": str(input_docx), "template": "report"})

    try:
        assert result["ok"] is True
        assert "output_docx_path" not in result
        assert "llm_analysis_path" not in result
        assert Path(result["markdown_report_path"]).exists()
        assert Path(result["json_detail_path"]).exists()
    finally:
        _cleanup_result_files(result)


def test_apply_docx_template_use_llm_writes_all_artifacts_and_keeps_operations_clean(
    tmp_path,
    monkeypatch,
):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    before = _docx_text_snapshot(input_docx)
    calls = []

    def fake_build_assistance(document, report):
        calls.append((document, report))
        return _fake_llm_assistance()

    monkeypatch.setattr(tools, "build_llm_assistance", fake_build_assistance)

    result = tools.apply_docx_template(
        {"input_path": str(input_docx), "template": "report", "use_llm": True}
    )

    try:
        assert result["ok"] is True
        assert len(calls) == 1
        assert result["llm_status"] == "ok"
        assert result["llm_mode"] == "private_llm"

        output_docx = Path(result["output_docx_path"])
        markdown_report = Path(result["markdown_report_path"])
        json_detail = Path(result["json_detail_path"])
        llm_analysis = Path(result["llm_analysis_path"])

        assert output_docx.exists()
        assert markdown_report.exists()
        assert json_detail.exists()
        assert llm_analysis.exists()
        assert "LLM 辅助分析" in markdown_report.read_text(encoding="utf-8")

        operations_text = json_detail.read_text(encoding="utf-8")
        operations_payload = json.loads(operations_text)
        assert "llm_assistance" not in operations_payload
        assert "template_recommendation" not in operations_text
        assert "llm-only-operation" not in operations_text
        assert _docx_text_snapshot(output_docx) == before
    finally:
        _cleanup_result_files(result)


def test_apply_docx_template_use_llm_exception_still_writes_formatted_docx_and_fallback(
    tmp_path,
    monkeypatch,
):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    def raise_llm_error(document, report):
        raise RuntimeError("llm exploded")

    monkeypatch.setattr(tools, "build_llm_assistance", raise_llm_error)

    result = tools.apply_docx_template(
        {"input_path": str(input_docx), "template": "report", "use_llm": True}
    )

    try:
        assert result["ok"] is True
        assert Path(result["output_docx_path"]).exists()
        assert Path(result["markdown_report_path"]).exists()
        assert Path(result["json_detail_path"]).exists()
        llm_analysis = Path(result["llm_analysis_path"])
        assert llm_analysis.exists()

        analysis = json.loads(llm_analysis.read_text(encoding="utf-8"))
        assert analysis["status"] == "fallback"
        assert analysis["mode"] == "rule_only_fallback"
        assert analysis["operations_source"] == "rule_engine_only"
        assert analysis["operations_generated"] is False
    finally:
        _cleanup_result_files(result)


def test_apply_docx_template_use_llm_false_does_not_call_llm(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    def fail_if_called(document, report):
        raise AssertionError("LLM should not be called when use_llm is false")

    monkeypatch.setattr(tools, "build_llm_assistance", fail_if_called)

    result = tools.apply_docx_template(
        {"input_path": str(input_docx), "template": "report", "use_llm": False}
    )

    try:
        assert result["ok"] is True
        assert "llm_analysis_path" not in result
    finally:
        _cleanup_result_files(result)


def test_generate_review_report_use_llm_writes_analysis_and_keeps_operations_clean(
    tmp_path,
    monkeypatch,
):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)
    calls = []

    def fake_build_assistance(document, report):
        calls.append((document, report))
        return _fake_llm_assistance()

    monkeypatch.setattr(tools, "build_llm_assistance", fake_build_assistance)

    result = tools.generate_review_report(
        {"input_path": str(input_docx), "template": "report", "use_llm": True}
    )

    try:
        assert result["ok"] is True
        assert len(calls) == 1
        assert "output_docx_path" not in result
        assert result["llm_status"] == "ok"
        assert result["llm_mode"] == "private_llm"

        markdown_report = Path(result["markdown_report_path"])
        json_detail = Path(result["json_detail_path"])
        llm_analysis = Path(result["llm_analysis_path"])

        assert markdown_report.exists()
        assert json_detail.exists()
        assert llm_analysis.exists()
        assert "LLM 辅助分析" in markdown_report.read_text(encoding="utf-8")

        operations_text = json_detail.read_text(encoding="utf-8")
        operations_payload = json.loads(operations_text)
        assert "llm_assistance" not in operations_payload
        assert "template_recommendation" not in operations_text
        assert "llm-only-operation" not in operations_text
    finally:
        _cleanup_result_files(result)


def test_unconnected_mcp_tools_do_not_call_llm_when_use_llm_is_present(tmp_path, monkeypatch):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    def fail_if_called(document, report):
        raise AssertionError("LLM should not be called by this MCP tool")

    monkeypatch.setattr(tools, "build_llm_assistance", fail_if_called)

    templates = tools.list_templates({"use_llm": True})
    analysis = tools.analyze_docx({"input_path": str(input_docx), "use_llm": True})
    check = tools.check_docx_style(
        {"input_path": str(input_docx), "template": "report", "use_llm": True}
    )

    assert templates["ok"] is True
    assert analysis["ok"] is True
    assert check["ok"] is True
    assert "llm_analysis_path" not in templates
    assert "llm_analysis_path" not in analysis
    assert "llm_analysis_path" not in check


def test_apply_docx_template_rejects_output_dir_outside_allowed_roots(tmp_path):
    input_docx = tmp_path / "input.docx"
    _create_docx(input_docx)

    result = tools.apply_docx_template(
        {
            "input_path": str(input_docx),
            "template": "report",
            "output_dir": str(tmp_path),
        }
    )

    assert result["ok"] is False
    assert result["error"]["type"] == "ValueError"
    assert "output_dir must be inside" in result["error"]["message"]
