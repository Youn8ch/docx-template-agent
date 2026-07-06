from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

import src.engine.formatter.apply_template as apply_template
from src.engine.checker.operation_validator import validate_operations
from src.engine.checker.style_checker import check_styles
from src.engine.formatter.apply_template import apply_operations_transactional
from src.engine.model.document_model import DocumentModel, ParagraphInfo, TableCellInfo, TableInfo
from src.engine.model.operation_model import FormatOperation, OperationResult, StyleCheckReport, StyleIssue
from src.engine.parser.docx_parser import parse_docx
from src.engine.parser.structure_detector import detect_structure
from src.engine.safety.content_integrity import content_snapshot


def _template() -> dict:
    return {
        "template_id": "test",
        "rules": {
            "title": {
                "font_name": "SimHei",
                "font_size": 22,
                "bold": True,
                "alignment": "center",
                "line_spacing": 1.5,
                "space_before": 0,
                "space_after": 12,
                "first_line_indent_chars": 0,
            },
            "heading_1": {
                "font_name": "SimHei",
                "font_size": 16,
                "bold": True,
                "alignment": "left",
                "line_spacing": 1.5,
                "space_before": 12,
                "space_after": 6,
                "first_line_indent_chars": 0,
            },
            "body": {
                "font_name": "SimSun",
                "font_size": 12,
                "bold": False,
                "alignment": "justify",
                "line_spacing": 1.5,
                "space_before": 0,
                "space_after": 0,
                "first_line_indent_chars": 2,
            },
            "table": {
                "font_name": "SimSun",
                "font_size": 10.5,
                "bold": False,
                "alignment": "center",
            },
        },
        "safety": {},
    }


def _create_docx(path: Path) -> None:
    document = Document()
    title = document.add_paragraph()
    title.paragraph_format.first_line_indent = Pt(24)
    title.add_run("Report Title")
    heading = document.add_paragraph()
    heading.paragraph_format.first_line_indent = Pt(24)
    heading.add_run("1. First Heading")
    body = document.add_paragraph()
    body.paragraph_format.first_line_indent = Pt(0)
    body.add_run("This body text must not change.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text must not change."
    document.save(path)


def _prepared_report(path: Path) -> tuple[dict, DocumentModel, StyleCheckReport]:
    template = _template()
    document = detect_structure(parse_docx(path))
    report = check_styles(document, template)
    return template, document, report


def _transaction(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)
    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
    )
    return source, output, template, document, report, execution


def test_first_line_indent_zero_is_idempotent_after_save_and_reopen(tmp_path):
    source, output, template, document, report, execution = _transaction(tmp_path)

    assert source.exists()
    assert output.exists()
    assert execution.status == "success"
    assert execution.issues_before == report.issue_count
    assert execution.issues_after == 0
    assert execution.operations_before == report.operation_count
    assert execution.operations_after == 0

    reparsed = parse_docx(output)
    roles = {paragraph.index: paragraph.role for paragraph in document.paragraphs}
    reparsed = reparsed.model_copy(
        update={
            "paragraphs": [
                paragraph.model_copy(update={"role": roles[paragraph.index]})
                for paragraph in reparsed.paragraphs
            ]
        }
    )
    after_report = check_styles(reparsed, template)

    assert after_report.issue_count == 0
    assert after_report.operation_count == 0


def test_title_zero_indent_is_written_and_rechecked_after_reopen(tmp_path):
    source, output, template, document, report, execution = _transaction(tmp_path)

    assert execution.status == "success"
    reopened = Document(str(output))
    heading_indent = reopened.paragraphs[0]._p.pPr.ind
    assert heading_indent.get(qn("w:firstLineChars")) == "0"
    assert heading_indent.get(qn("w:firstLine")) is None
    assert heading_indent.get(qn("w:hanging")) is None
    assert heading_indent.get(qn("w:hangingChars")) is None

    reparsed = parse_docx(output)
    heading = reparsed.paragraphs[0].model_copy(update={"role": "title"})
    assert heading.first_line_indent_chars == 0
    checked = check_styles(
        reparsed.model_copy(update={"paragraphs": [heading]}),
        {"template_id": "indent", "rules": {"title": {"first_line_indent_chars": 0}}},
    )
    assert checked.operation_count == 0


def test_zero_indent_normalization_is_limited_to_non_indented_roles():
    template = {"template_id": "indent", "rules": {"body": {"first_line_indent_chars": 0}}}
    document = DocumentModel(
        filepath=Path("sample.docx"),
        paragraphs=[
            ParagraphInfo(
                index=1,
                text="body",
                role="body",
                first_line_indent=None,
                first_line_indent_chars=None,
            )
        ],
    )

    report = check_styles(document, template)

    assert report.issue_count == 1
    assert report.operations[0].properties["first_line_indent_chars"] == 0


@pytest.mark.parametrize("role", ["body", "heading_1"])
def test_hanging_conflict_is_not_treated_as_zero_indent(role):
    template = {"template_id": "indent", "rules": {role: {"first_line_indent_chars": 0}}}
    document = DocumentModel(
        filepath=Path("sample.docx"),
        paragraphs=[
            ParagraphInfo(
                index=1,
                text="paragraph",
                role=role,
                first_line_indent=None,
                first_line_indent_chars=None,
                hanging=240,
                hanging_chars=None,
            )
        ],
    )

    report = check_styles(document, template)

    assert report.issue_count == 1


def test_heading_explicit_zero_and_missing_zero_indent_match():
    template = {"template_id": "indent", "rules": {"heading_1": {"first_line_indent_chars": 0}}}
    document = DocumentModel(
        filepath=Path("sample.docx"),
        paragraphs=[
            ParagraphInfo(index=1, text="explicit", role="heading_1", first_line_indent_chars=0),
            ParagraphInfo(index=2, text="missing", role="heading_1", first_line_indent_chars=None),
        ],
    )

    report = check_styles(document, template)

    assert report.issue_count == 0
    assert report.operation_count == 0


def test_transaction_rechecks_by_reopening_saved_docx(tmp_path, monkeypatch):
    calls = {"count": 0}
    real_parse = apply_template.parse_docx

    def counting_parse(path):
        calls["count"] += 1
        assert Path(path).exists()
        return real_parse(path)

    monkeypatch.setattr(apply_template, "parse_docx", counting_parse)

    *_rest, execution = _transaction(tmp_path)

    assert execution.status == "success"
    assert calls["count"] == 1


def test_content_integrity_detects_paragraph_text_change(tmp_path, monkeypatch):
    real_execute = apply_template._execute_operations

    def mutating_execute(document, operations):
        results = real_execute(document, operations)
        document.paragraphs[0].text = "Changed title"
        return results

    monkeypatch.setattr(apply_template, "_execute_operations", mutating_execute)
    *_rest, output, _template_value, _document, _report, execution = _transaction(tmp_path)

    assert execution.status == "content_integrity_failed"
    assert "paragraph text changed" in execution.integrity_errors
    assert not output.exists()
    assert execution.temp_file_retained is False
    assert execution.temp_output_path is None


def test_transaction_can_retain_failed_temp_for_cli_debugging(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)
    real_execute = apply_template._execute_operations

    def mutating_execute(docx, operations):
        results = real_execute(docx, operations)
        docx.paragraphs[0].text = "Changed title"
        return results

    monkeypatch.setattr(apply_template, "_execute_operations", mutating_execute)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
        retain_failed_temp=True,
    )

    assert execution.status == "content_integrity_failed"
    assert execution.temp_file_retained is True
    assert execution.temp_output_path
    assert Path(execution.temp_output_path).exists()
    assert Path(execution.temp_output_path).parent == output.parent
    assert not output.exists()


def test_success_removes_temp_and_records_output_overwrite(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    output.write_bytes(b"old output")
    template, document, report = _prepared_report(source)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
    )

    assert execution.status == "success"
    assert execution.output_path == str(output)
    assert execution.output_existed_before is True
    assert execution.output_overwritten is True
    assert execution.temp_file_retained is False
    assert execution.temp_output_path is None
    assert output.exists()


def test_existing_output_can_be_rejected_without_overwrite(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    output.write_bytes(b"old output")
    before = output.read_bytes()
    template, document, report = _prepared_report(source)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
        overwrite_output=False,
    )

    assert execution.status == "validation_failed"
    assert execution.output_path is None
    assert execution.expected_output_path == str(output)
    assert execution.output_existed_before is True
    assert output.read_bytes() == before


def test_save_exception_deletes_failed_temp_by_default(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)
    real_document = apply_template.Document

    class FailingSaveDocument:
        def __init__(self, path):
            self._document = real_document(path)

        def __getattr__(self, name):
            return getattr(self._document, name)

        def save(self, path):
            raise RuntimeError("save failed")

    monkeypatch.setattr(apply_template, "Document", FailingSaveDocument)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
    )

    assert execution.status == "fatal"
    assert execution.temp_file_retained is False
    assert execution.temp_output_path is None
    assert not output.exists()


def test_recheck_exception_deletes_failed_temp_by_default(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)

    def failing_parse(path):
        raise RuntimeError("recheck failed")

    monkeypatch.setattr(apply_template, "parse_docx", failing_parse)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
    )

    assert execution.status == "fatal"
    assert execution.temp_file_retained is False
    assert execution.temp_output_path is None
    assert not output.exists()


def test_content_integrity_detects_table_text_change(tmp_path, monkeypatch):
    real_execute = apply_template._execute_operations

    def mutating_execute(document, operations):
        results = real_execute(document, operations)
        document.tables[0].cell(0, 0).text = "Changed table text"
        return results

    monkeypatch.setattr(apply_template, "_execute_operations", mutating_execute)
    *_rest, output, _template_value, _document, _report, execution = _transaction(tmp_path)

    assert execution.status == "content_integrity_failed"
    assert "table cell text changed" in execution.integrity_errors
    assert not output.exists()


def test_content_integrity_detects_paragraph_count_change(tmp_path, monkeypatch):
    real_execute = apply_template._execute_operations

    def mutating_execute(document, operations):
        results = real_execute(document, operations)
        document.add_paragraph("Unexpected paragraph")
        return results

    monkeypatch.setattr(apply_template, "_execute_operations", mutating_execute)
    *_rest, output, _template_value, _document, _report, execution = _transaction(tmp_path)

    assert execution.status == "content_integrity_failed"
    assert "paragraph count changed" in execution.integrity_errors
    assert "paragraph index structure changed" in execution.recheck_errors
    assert not output.exists()


def test_content_integrity_detects_table_count_change(tmp_path, monkeypatch):
    real_execute = apply_template._execute_operations

    def mutating_execute(document, operations):
        results = real_execute(document, operations)
        document.add_table(rows=1, cols=1)
        return results

    monkeypatch.setattr(apply_template, "_execute_operations", mutating_execute)
    *_rest, output, _template_value, _document, _report, execution = _transaction(tmp_path)

    assert execution.status == "content_integrity_failed"
    assert "table count changed" in execution.integrity_errors
    assert "table count changed before recheck" in execution.recheck_errors
    assert not output.exists()


def test_content_snapshot_uses_structured_paragraph_boundaries():
    first = DocumentModel(
        filepath=Path("a.docx"),
        paragraphs=[
            ParagraphInfo(index=1, text="ab"),
            ParagraphInfo(index=2, text="c"),
            ParagraphInfo(index=3, text=""),
        ],
    )
    second = DocumentModel(
        filepath=Path("b.docx"),
        paragraphs=[
            ParagraphInfo(index=1, text="a"),
            ParagraphInfo(index=2, text="bc"),
            ParagraphInfo(index=3, text=""),
        ],
    )

    first_snapshot = content_snapshot(first)
    second_snapshot = content_snapshot(second)

    assert first_snapshot["protection_scope"]["body_paragraphs"] is True
    assert first_snapshot["protection_scope"]["headers_footers_protected"] is False
    assert first_snapshot["paragraph_payload_fingerprint"] != second_snapshot[
        "paragraph_payload_fingerprint"
    ]


def test_content_snapshot_preserves_table_cell_paragraph_boundaries():
    one_cell_two_paragraphs = DocumentModel(
        filepath=Path("a.docx"),
        tables=[
            TableInfo(
                index=1,
                rows=1,
                cols=1,
                cells=[
                    TableCellInfo(
                        row_index=1,
                        col_index=1,
                        text="ab\nc",
                        paragraph_texts=["ab", "c"],
                    )
                ],
            )
        ],
    )
    one_cell_one_paragraph = DocumentModel(
        filepath=Path("b.docx"),
        tables=[
            TableInfo(
                index=1,
                rows=1,
                cols=1,
                cells=[
                    TableCellInfo(
                        row_index=1,
                        col_index=1,
                        text="ab\nc",
                        paragraph_texts=["ab\nc"],
                    )
                ],
            )
        ],
    )

    assert content_snapshot(one_cell_two_paragraphs)["table_payload_fingerprint"] != content_snapshot(
        one_cell_one_paragraph
    )["table_payload_fingerprint"]


def test_merged_cell_snapshot_records_current_python_docx_semantics(tmp_path):
    source = tmp_path / "merged.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "merged text"
    table.cell(0, 0).merge(table.cell(0, 1))
    document.save(source)

    parsed = parse_docx(source)
    snapshot_a = content_snapshot(parsed)
    snapshot_b = content_snapshot(parse_docx(source))

    assert parsed.table_count == 1
    assert parsed.tables[0].rows == 1
    assert parsed.tables[0].cols == 2
    assert len(parsed.tables[0].cells) == 2
    assert snapshot_a["table_payload_fingerprint"] == snapshot_b["table_payload_fingerprint"]


def test_operation_validation_rejects_illegal_property():
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[1],
        role="body",
        properties={"text": "mutate"},
    )

    errors = validate_operations([operation], document, _template())

    assert any("forbidden mutation field" in error for error in errors)


def test_operation_validation_rejects_unknown_property():
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[1],
        role="body",
        properties={"font_color": "red"},
    )

    errors = validate_operations([operation], document, _template())

    assert any("unsupported property" in error for error in errors)


def test_operation_validation_rejects_out_of_range_target():
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[2],
        role="body",
        properties={"font_size": 12},
    )

    errors = validate_operations([operation], document, _template())

    assert any("paragraph target out of range" in error for error in errors)


def test_operation_validation_rejects_illegal_numeric_value():
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[1],
        role="body",
        properties={"font_size": 500},
    )

    errors = validate_operations([operation], document, _template())

    assert any("font_size out of range" in error for error in errors)


@pytest.mark.parametrize(
    ("field", "value", "reason"),
    [
        ("font_size", True, "boolean is not numeric"),
        ("line_spacing", False, "boolean is not numeric"),
        ("font_size", float("nan"), "value must be finite"),
        ("font_size", "NaN", "value must be finite"),
        ("font_size", float("inf"), "value must be finite"),
        ("font_size", "-Infinity", "value must be finite"),
        ("font_size", "not-a-number", "value is not numeric"),
    ],
)
def test_operation_validation_rejects_bool_and_non_finite_numbers(field, value, reason):
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[1],
        role="body",
        properties={field: value},
    )

    errors = validate_operations([operation], document, _template())

    assert any(field in error and reason in error for error in errors)


@pytest.mark.parametrize(
    ("field", "value"),
    [
        ("font_size", 12),
        ("line_spacing", 1.5),
        ("space_before", "6"),
        ("first_line_indent_chars", "2"),
    ],
)
def test_operation_validation_accepts_finite_numbers(field, value):
    document = DocumentModel(filepath=Path("x.docx"), paragraphs=[ParagraphInfo(index=1, role="body")])
    operation = FormatOperation(
        operation_id="op-1",
        action="apply_paragraph_style",
        target_type="paragraph",
        target_indices=[1],
        role="body",
        properties={field: value},
    )

    assert validate_operations([operation], document, _template()) == []


def test_format_operation_rejects_unknown_top_level_field():
    with pytest.raises(Exception):
        FormatOperation(
            operation_id="op-1",
            action="apply_paragraph_style",
            target_type="paragraph",
            target_indices=[1],
            properties={"font_size": 12},
            unexpected=True,
        )


def test_transaction_validation_failure_does_not_write_formal_output(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)
    invalid = [
        FormatOperation(
            operation_id="op-bad",
            action="apply_paragraph_style",
            target_type="paragraph",
            target_indices=[999],
            role="body",
            properties={"font_size": 12},
        )
    ]

    execution = apply_operations_transactional(
        source,
        output,
        invalid,
        template=template,
        document_before=document,
        report_before=report,
    )

    assert execution.status == "validation_failed"
    assert execution.validation_errors
    assert not output.exists()


def test_recheck_failure_keeps_temp_and_does_not_write_formal_output(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _create_docx(source)
    template, document, report = _prepared_report(source)

    def failing_recheck(document_after, template_after):
        return StyleCheckReport(
            template_id="test",
            issues=[
                StyleIssue(
                    issue_id="issue-after",
                    issue_type="font_size",
                    target_type="paragraph",
                    target_index=1,
                    role="title",
                    current=10,
                    expected=22,
                    message="forced recheck issue",
                )
            ],
        )

    monkeypatch.setattr(apply_template, "check_styles", failing_recheck)

    execution = apply_operations_transactional(
        source,
        output,
        report.operations,
        template=template,
        document_before=document,
        report_before=report,
    )

    assert execution.status == "validation_failed"
    assert execution.issues_after == 1
    assert execution.operations_after == 0
    assert not output.exists()
    assert execution.temp_file_retained is False
    assert execution.temp_output_path is None
