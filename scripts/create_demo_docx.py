"""Create a demo docx file for offline formatting tests."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = Path(__file__).resolve().parents[1] / "samples" / "input" / "demo.docx"


def set_run_font(run, font_name: str, font_size_pt: float, bold: bool | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.font.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(
    document: Document,
    text: str,
    font_name: str,
    font_size_pt: float,
    first_line_indent_pt: float | None = None,
    bold: bool | None = None,
):
    paragraph = document.add_paragraph()
    if first_line_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size_pt, bold)
    return paragraph


def create_demo_docx(output_path: Path = OUTPUT_PATH) -> Path:
    document = Document()

    add_paragraph(document, "项目建设方案", "宋体", 18, bold=True)
    add_paragraph(document, "一、项目背景", "宋体", 16, bold=True)
    add_paragraph(document, "（一）建设现状", "宋体", 14, bold=True)
    add_paragraph(document, "1. 当前问题", "宋体", 12, bold=True)

    body_paragraphs = [
        ("当前项目文档来源较多，格式标准不完全一致，影响后续归档和审核效率。", 21),
        ("部分材料在段落缩进、字体字号和标题层级方面存在差异，需要进行统一排版。", None),
        ("本示例文档用于验证离线 docx 模板化处理流程，不包含真实业务数据。", 21),
        ("处理过程应保持正文内容不变，仅对明确允许的格式属性进行调整。", None),
    ]
    for text, first_line_indent in body_paragraphs:
        add_paragraph(document, text, "微软雅黑", 10.5, first_line_indent)

    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table_data = [
        ["序号", "事项", "说明"],
        ["1", "资料收集", "汇总现有项目资料"],
        ["2", "格式检查", "识别不符合模板的格式"],
        ["3", "输出结果", "生成新的排版后文档"],
    ]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(table_data[row_index][col_index])
            set_run_font(run, "宋体", 10.5, bold=(row_index == 0))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    created_path = create_demo_docx()
    print(f"created demo docx: {created_path}")
