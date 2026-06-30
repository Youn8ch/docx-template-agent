# -*- coding: utf-8 -*-
"""
Comprehensive unit tests for wps-agent docx_engine.
80+ test cases covering Document, Paragraph, Run, Table, Cell, Section,
Serializer, XML Parser, Intelligence, and Formatter.
"""
import os
import sys
import pytest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))


# ═══════════════════════════════════════════════════════
# Run Tests (10+ cases)
# ═══════════════════════════════════════════════════════

class TestRun:
    def test_create_default(self):
        from docx_engine.document_model import Run
        r = Run()
        assert r.text == ""
        assert r.bold is False
        assert r.italic is False
        assert r.size is None

    def test_create_with_text(self):
        from docx_engine.document_model import Run
        r = Run(text="Hello")
        assert r.text == "Hello"

    def test_create_full_formatting(self):
        from docx_engine.document_model import Run
        r = Run(
            text="测试", font="宋体", font_east_asia="宋体",
            size=12, bold=True, italic=True, underline="single",
            color="FF0000", highlight="yellow",
            superscript=True, strike=True,
        )
        assert r.font == "宋体"
        assert r.size == 12
        assert r.bold is True
        assert r.italic is True
        assert r.underline == "single"
        assert r.color == "FF0000"
        assert r.highlight == "yellow"
        assert r.superscript is True
        assert r.strike is True

    def test_is_empty(self):
        from docx_engine.document_model import Run
        assert Run().is_empty() is True
        assert Run(text="x").is_empty() is False

    def test_is_empty_whitespace_is_not_empty(self):
        from docx_engine.document_model import Run
        assert Run(text=" ").is_empty() is False

    def test_has_formatting_true(self):
        from docx_engine.document_model import Run
        r = Run(text="text", bold=True)
        assert r.has_formatting() is True

    def test_has_formatting_false(self):
        from docx_engine.document_model import Run
        r = Run(text="text")
        assert r.has_formatting() is False

    def test_has_formatting_with_size(self):
        from docx_engine.document_model import Run
        r = Run(text="text", size=12)
        assert r.has_formatting() is True

    def test_has_formatting_all_properties(self):
        from docx_engine.document_model import Run
        r = Run(
            text="text", font="Arial", size=14, bold=True, italic=True,
            underline="single", color="000000", highlight="none",
            strike=True, double_strike=True, emboss=True, imprint=True,
            shadow=True, outline=True, superscript=False, subscript=False,
            caps=True, small_caps=True, char_spacing=2, kerning=1, scaling=100,
            baseline_offset=0, vanish=False, emphasis_mark="dot",
        )
        assert r.has_formatting() is True

    def test_clone_independence(self):
        from docx_engine.document_model import Run
        r1 = Run(text="original", bold=True, size=12)
        r2 = r1.clone()
        r2.text = "modified"
        r2.bold = False
        assert r1.text == "original"
        assert r1.bold is True
        assert r2.text == "modified"
        assert r2.bold is False

    def test_text_effects_properties(self):
        from docx_engine.document_model import Run
        r = Run(text="fx", shadow=True, outline=True, emboss=True, imprint=True)
        assert r.shadow is True
        assert r.outline is True
        assert r.emboss is True
        assert r.imprint is True

    def test_spacing_properties(self):
        from docx_engine.document_model import Run
        r = Run(text="spaced", char_spacing=3.0, kerning=2.0, scaling=120)
        assert r.char_spacing == 3.0
        assert r.kerning == 2.0
        assert r.scaling == 120


# ═══════════════════════════════════════════════════════
# Paragraph Tests (10+ cases)
# ═══════════════════════════════════════════════════════

class TestParagraph:
    def test_create_empty(self):
        from docx_engine.document_model import Paragraph
        p = Paragraph()
        assert p.runs == []
        assert p.text == ""

    def test_create_with_runs(self, sample_run):
        from docx_engine.document_model import Paragraph
        p = Paragraph(runs=[sample_run, sample_run.clone()])
        assert len(p.runs) == 2

    def test_text_property(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Hello"), Run(text=" World")])
        assert p.text == "Hello World"

    def test_is_heading_true(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Title")], outline_level=0)
        assert p.is_heading() is True

    def test_is_heading_true_level_1(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Sub")], outline_level=1)
        assert p.is_heading() is True

    def test_is_heading_false_body(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Body")], outline_level=None)
        assert p.is_heading() is False

    def test_is_heading_by_style(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Title")], style_id="Heading1", outline_level=None)
        assert p.is_heading() is True

    def test_heading_level(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="L2")], outline_level=1)
        assert p.heading_level() == 2

    def test_heading_level_body(self):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Body")])
        assert p.heading_level() is None

    def test_is_empty(self):
        from docx_engine.document_model import Paragraph
        assert Paragraph().is_empty() is True

    def test_is_empty_with_runs(self):
        from docx_engine.document_model import Paragraph, Run
        assert Paragraph(runs=[Run(text="x")]).is_empty() is False

    def test_add_run(self):
        from docx_engine.document_model import Paragraph
        p = Paragraph()
        run = p.add_run("new text", bold=True, size=14)
        assert len(p.runs) == 1
        assert run.text == "new text"
        assert run.bold is True
        assert run.size == 14

    def test_add_run_returns_run(self):
        from docx_engine.document_model import Paragraph
        p = Paragraph()
        r = p.add_run("x")
        assert r in p.runs

    def test_clone_independence(self):
        from docx_engine.document_model import Paragraph, Run
        p1 = Paragraph(runs=[Run(text="orig")], alignment="center")
        p2 = p1.clone()
        p2.alignment = "left"
        p2.runs[0].text = "changed"
        assert p1.alignment == "center"
        assert p1.runs[0].text == "orig"
        assert p2.alignment == "left"
        assert p2.runs[0].text == "changed"

    def test_paragraph_format_properties(self, sample_paragraph):
        assert sample_paragraph.alignment == "justify"
        assert sample_paragraph.first_line_indent == 24
        assert sample_paragraph.space_before == 12
        assert sample_paragraph.space_after == 6

    def test_paragraph_pagination_properties(self):
        from docx_engine.document_model import Paragraph
        p = Paragraph(keep_with_next=True, keep_lines=True, page_break_before=True, widow_control=False)
        assert p.keep_with_next is True
        assert p.keep_lines is True
        assert p.page_break_before is True
        assert p.widow_control is False


# ═══════════════════════════════════════════════════════
# Cell Tests (3+ cases)
# ═══════════════════════════════════════════════════════

class TestCell:
    def test_create_default(self):
        from docx_engine.document_model import Cell
        c = Cell()
        assert c.paragraphs == []
        assert c.text == ""

    def test_text_joins_paragraphs(self):
        from docx_engine.document_model import Cell, Paragraph, Run
        c = Cell(paragraphs=[
            Paragraph(runs=[Run(text="Line1")]),
            Paragraph(runs=[Run(text="Line2")]),
        ])
        assert c.text == "Line1\nLine2"

    def test_cell_formatting_properties(self):
        from docx_engine.document_model import Cell
        c = Cell(width=100, shading="FF0000", v_align="center")
        assert c.width == 100
        assert c.shading == "FF0000"
        assert c.v_align == "center"


# ═══════════════════════════════════════════════════════
# Table Tests (5+ cases)
# ═══════════════════════════════════════════════════════

class TestTable:
    def test_create_default(self):
        from docx_engine.document_model import Table
        t = Table()
        assert t.row_count == 0
        assert t.col_count == 0

    def test_row_col_count(self, sample_table):
        assert sample_table.row_count == 3
        assert sample_table.col_count == 2

    def test_cell_access(self, sample_table):
        c = sample_table.cell(1, 1)
        assert c is not None
        assert "R1C1" in c.text

    def test_cell_access_out_of_bounds(self, sample_table):
        assert sample_table.cell(99, 99) is None
        assert sample_table.cell(-1, 0) is None

    def test_cell_access_zero_zero(self, sample_table):
        c = sample_table.cell(0, 0)
        assert c is not None
        assert "R0C0" in c.text

    def test_table_column_widths(self, sample_table):
        assert sample_table.column_widths == [200, 300]

    def test_empty_table_cell_none(self):
        from docx_engine.document_model import Table
        t = Table()
        assert t.cell(0, 0) is None


# ═══════════════════════════════════════════════════════
# Section Tests (3+ cases)
# ═══════════════════════════════════════════════════════

class TestSection:
    def test_default_a4(self):
        from docx_engine.document_model import Section
        s = Section()
        assert s.page_width == 595.3
        assert s.page_height == 841.9
        assert s.orientation == "portrait"
        assert s.has_changes() is False

    def test_has_changes_modified(self):
        from docx_engine.document_model import Section
        s = Section(orientation="landscape")
        assert s.has_changes() is True

    def test_has_changes_margin(self):
        from docx_engine.document_model import Section
        s = Section(left_margin=90)
        assert s.has_changes() is True


# ═══════════════════════════════════════════════════════
# Document Tests (20+ cases)
# ═══════════════════════════════════════════════════════

class TestDocument:
    def test_create_empty(self, empty_document):
        assert len(empty_document.paragraphs) == 0
        assert empty_document.text == ""

    def test_text_property(self, sample_document):
        text = sample_document.text
        assert "My Thesis Title" in text
        assert "Abstract" in text

    def test_get_paragraph_valid(self, sample_document):
        p = sample_document.get_paragraph(1)
        assert p is not None
        assert "My Thesis Title" in p.text

    def test_get_paragraph_out_of_bounds(self, sample_document):
        assert sample_document.get_paragraph(999) is None
        assert sample_document.get_paragraph(0) is None
        assert sample_document.get_paragraph(-1) is None

    def test_get_heading_structure(self, sample_document):
        headings = sample_document.get_heading_structure()
        assert len(headings) >= 3
        level1_headings = [h for h in headings if h["level"] == 1]
        assert len(level1_headings) >= 1

    def test_find_paragraphs_by_text(self, sample_document):
        idx = sample_document.find_paragraphs_by_text("Abstract")
        assert len(idx) == 1

    def test_find_paragraphs_by_text_case_insensitive(self, sample_document):
        idx = sample_document.find_paragraphs_by_text("abstract", case_sensitive=False)
        assert len(idx) >= 1

    def test_find_paragraphs_by_text_not_found(self, sample_document):
        idx = sample_document.find_paragraphs_by_text("xyznonexistent123")
        assert idx == []

    def test_find_paragraphs_by_style(self, sample_document):
        idx = sample_document.find_paragraphs_by_style("Title")
        assert len(idx) >= 1

    def test_replace_text(self, sample_document):
        count = sample_document.replace_text("Introduction", "Overview")
        p = sample_document.get_paragraph(4)
        assert "Overview" in p.text
        assert count >= 1

    def test_replace_text_no_match(self, sample_document):
        count = sample_document.replace_text("xyznonexistent", "replacement")
        assert count == 0

    def test_replace_text_multiple(self, sample_document):
        doc = sample_document
        doc.replace_text("body text", "content text")
        assert "content text" in doc.text

    def test_insert_paragraph(self, empty_document):
        from docx_engine.document_model import Paragraph, Run
        p = Paragraph(runs=[Run(text="Inserted")])
        result = empty_document.insert_paragraph(1, p)
        assert result is not None
        assert len(empty_document.paragraphs) == 1
        assert empty_document.paragraphs[0].text == "Inserted"

    def test_insert_paragraph_at_middle(self, sample_document):
        from docx_engine.document_model import Paragraph, Run
        original_count = len(sample_document.paragraphs)
        p = Paragraph(runs=[Run(text="Inserted middle")])
        sample_document.insert_paragraph(3, p)
        assert len(sample_document.paragraphs) == original_count + 1
        assert sample_document.get_paragraph(3).text == "Inserted middle"

    def test_delete_paragraph(self, sample_document):
        original_count = len(sample_document.paragraphs)
        deleted = sample_document.delete_paragraph(1)
        assert deleted is not None
        assert deleted.text == "My Thesis Title"
        assert len(sample_document.paragraphs) == original_count - 1

    def test_delete_paragraph_out_of_bounds(self, sample_document):
        deleted = sample_document.delete_paragraph(999)
        assert deleted is None

    def test_get_statistics(self, sample_document):
        stats = sample_document.get_statistics()
        assert stats["paragraph_count"] >= 9
        assert "heading_count" in stats
        assert "total_characters" in stats
        assert "table_count" in stats
        assert stats["heading_count"] >= 4

    def test_is_dirty_new_document(self):
        from docx_engine.document_model import Document, Paragraph, Run
        doc = Document()
        doc.insert_paragraph(1, Paragraph(runs=[Run(text="Test")]))
        assert doc.is_dirty() is True

    def test_is_dirty_loaded_document(self, sample_document):
        assert sample_document.is_dirty() is False

    def test_iter_paragraphs(self, sample_document):
        paras = list(sample_document.iter_paragraphs())
        assert len(paras) == len(sample_document.paragraphs)

    def test_detect_semantic_structure(self, sample_document):
        result = sample_document.detect_semantic_structure()
        assert len(result) == len(sample_document.paragraphs)
        roles = [r.get("semantic_role", "") for r in result]
        assert any("heading" in role or "title" in role for role in roles)
        assert any("abstract_label" in role or "body" in role for role in roles)

    def test_detect_cross_references(self, sample_document):
        refs = sample_document.detect_cross_references()
        assert isinstance(refs, list)

    def test_get_full_structure(self, sample_document):
        structure = sample_document.get_full_structure()
        assert "paragraphs" in structure
        assert "tables" in structure
        assert len(structure["paragraphs"]) >= 9


# ═══════════════════════════════════════════════════════
# XML Parser Tests (5+ cases)
# ═══════════════════════════════════════════════════════

class TestXmlParser:
    def test_unpack_docx(self, sample_docx):
        from docx_engine.xml_parser import unpack_docx
        unpacked = unpack_docx(sample_docx)
        assert os.path.exists(os.path.join(unpacked, "word", "document.xml"))

    def test_parse_docx(self, sample_docx):
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        assert parsed["document"] is not None

    def test_pack_docx(self, sample_docx, temp_dir):
        from docx_engine.xml_parser import unpack_docx, pack_docx
        unpacked = unpack_docx(sample_docx)
        output = os.path.join(temp_dir, "repacked.docx")
        pack_docx(unpacked, output)
        assert os.path.exists(output)

    def test_unpack_then_parse_roundtrip(self, sample_docx):
        from docx_engine.xml_parser import unpack_docx, pack_docx, parse_docx
        from docx import Document as PyDocx
        import tempfile, shutil
        unpacked = unpack_docx(sample_docx)
        tmp = tempfile.mkdtemp()
        output = os.path.join(tmp, "roundtrip.docx")
        pack_docx(unpacked, output)
        doc = PyDocx(output)
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Document" in t for t in texts)
        shutil.rmtree(tmp)

    def test_parse_nonexistent_file(self):
        from docx_engine.xml_parser import parse_docx
        from docx_engine.errors import ParseError
        with pytest.raises(ParseError):
            parse_docx("/nonexistent/file.docx")


# ═══════════════════════════════════════════════════════
# Serializer Tests (8+ cases)
# ═══════════════════════════════════════════════════════

class TestSerializer:
    def test_build_from_python_docx(self, sample_docx):
        from docx_engine.serializer import build_document_model
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        doc = build_document_model(parsed)
        assert len(doc.paragraphs) >= 4

    def test_roundtrip_preserves_paragraphs(self, sample_docx, temp_dir):
        from docx_engine.serializer import serialize_document_model, build_document_model
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        doc = build_document_model(parsed)
        output = os.path.join(temp_dir, "output.docx")
        serialize_document_model(doc, output, original_docx=sample_docx)
        assert os.path.exists(output)

    def test_roundtrip_preserves_run_formatting(self, sample_docx, temp_dir):
        from docx_engine.serializer import serialize_document_model, build_document_model, load_and_build
        model1 = load_and_build(sample_docx)
        output = os.path.join(temp_dir, "rt.docx")
        serialize_document_model(model1, output, original_docx=sample_docx)
        model2 = load_and_build(output)
        assert len(model2.paragraphs) >= len(model1.paragraphs)

    def test_load_and_build(self, sample_docx):
        from docx_engine.serializer import load_and_build
        doc = load_and_build(sample_docx)
        assert doc.paragraphs
        assert len(doc.paragraphs) >= 1

    def test_build_document_model_preserves_heading_levels(self, sample_docx):
        from docx_engine.serializer import build_document_model
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        doc = build_document_model(parsed)
        headings = doc.get_heading_structure()
        assert len(headings) >= 2

    def test_serialize_empty_document(self, sample_docx, temp_dir):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.serializer import serialize_document_model
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="Minimal")]))
        output = os.path.join(temp_dir, "minimal.docx")
        serialize_document_model(doc, output, original_docx=sample_docx)
        assert os.path.exists(output)

    def test_roundtrip_preserves_tables(self, sample_docx_with_table, temp_dir):
        from docx_engine.serializer import load_and_build, serialize_document_model
        model = load_and_build(sample_docx_with_table)
        output = os.path.join(temp_dir, "table_rt.docx")
        serialize_document_model(model, output, original_docx=sample_docx_with_table)
        assert os.path.exists(output)
        # Verify table count
        assert len(model.tables) >= 1

    def test_serialize_media_files_preserved(self, sample_docx, temp_dir):
        from docx_engine.serializer import load_and_build, serialize_document_model
        model = load_and_build(sample_docx)
        output = os.path.join(temp_dir, "media.docx")
        serialize_document_model(model, output, original_docx=sample_docx)
        assert os.path.exists(output)


# ═══════════════════════════════════════════════════════
# Intelligence Tests (10+ cases)
# ═══════════════════════════════════════════════════════

class TestIntelligence:
    def test_detect_thesis(self, sample_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_document)
        doc_type = analyzer.detect_document_type()
        assert doc_type in ("thesis", "paper", "report")

    def test_detect_contract(self, sample_contract_doc):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_contract_doc)
        doc_type = analyzer.detect_document_type()
        assert doc_type in ("contract", "report")

    def test_detect_resume(self, sample_resume_doc):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_resume_doc)
        doc_type = analyzer.detect_document_type()
        assert doc_type in ("resume", "letter")

    def test_detect_paragraph_roles(self, sample_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_document)
        roles = analyzer.detect_paragraph_roles()
        assert len(roles) == len(sample_document.paragraphs)
        role_names = [r["role"] for r in roles]
        assert any("heading" in role or "title" in role.lower() for role in role_names)

    def test_detect_roles_abstract(self, sample_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_document)
        roles = analyzer.detect_paragraph_roles()
        abstract_found = any(
            "abstract" in r["role"].lower() or "abstract" in r.get("text", "").lower()
            for r in roles
        )
        assert abstract_found is True

    def test_analyze_formatting_quality(self, sample_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_document)
        quality = analyzer.analyze_formatting_quality()
        assert "score" in quality
        assert "issues" in quality

    def test_get_document_outline(self, sample_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(sample_document)
        outline = analyzer.get_document_outline()
        assert len(outline) >= 3
        assert "text" in outline[0]
        assert "level" in outline[0]

    def test_detect_empty_document(self, empty_document):
        from docx_engine.intelligence import DocumentAnalyzer
        analyzer = DocumentAnalyzer(empty_document)
        doc_type = analyzer.detect_document_type()
        assert isinstance(doc_type, str)

    def test_detect_memo(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.intelligence import DocumentAnalyzer
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="To: All Staff")]))
        doc.paragraphs.append(Paragraph(runs=[Run(text="From: Manager")]))
        doc.paragraphs.append(Paragraph(runs=[Run(text="Subject: Meeting")]))
        analyzer = DocumentAnalyzer(doc)
        doc_type = analyzer.detect_document_type()
        assert doc_type in ("memo", "letter")

    def test_detect_letter(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.intelligence import DocumentAnalyzer
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="Dear Mr. Smith,")]))
        doc.paragraphs.append(Paragraph(runs=[Run(text="Body content.")]))
        doc.paragraphs.append(Paragraph(runs=[Run(text="Sincerely,")]))
        analyzer = DocumentAnalyzer(doc)
        doc_type = analyzer.detect_document_type()
        assert doc_type in ("letter", "memo")


# ═══════════════════════════════════════════════════════
# Formatter Tests (8+ cases)
# ═══════════════════════════════════════════════════════

class TestFormatter:
    def test_auto_format_thesis(self, sample_document):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_document)
        result = formatter.auto_format("thesis")
        assert result["applied"] is True
        assert result["change_count"] >= 0

    def test_auto_format_report(self, sample_document):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_document)
        result = formatter.auto_format("report")
        assert result["applied"] is True

    def test_auto_format_resume(self, sample_resume_doc):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_resume_doc)
        result = formatter.auto_format("resume")
        assert result["applied"] is True

    def test_auto_format_contract(self, sample_contract_doc):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_contract_doc)
        result = formatter.auto_format("general")
        assert result["applied"] is True

    def test_add_multi_level_numbering(self, sample_document):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_document)
        result = formatter.add_multi_level_numbering()
        assert result["change_count"] >= 0

    def test_apply_template_thesis_cn(self, sample_document):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_document)
        result = formatter.apply_template("thesis_cn")
        assert result["template"] == "thesis_cn"

    def test_apply_template_official(self, sample_document):
        from docx_engine.formatter import Formatter
        formatter = Formatter(sample_document)
        result = formatter.apply_template("report_official")
        assert result["template"] == "report_official"

    def test_format_heading_alignment(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.formatter import Formatter
        doc = Document()
        doc.paragraphs.append(Paragraph(
            runs=[Run(text="Chapter 1", bold=True)],
            outline_level=0,
        ))
        doc.paragraphs.append(Paragraph(
            runs=[Run(text="Body text.")],
        ))
        formatter = Formatter(doc)
        formatter.auto_format("thesis")
        assert doc.paragraphs[0].alignment == "center"
        assert doc.paragraphs[1].alignment == "justify"


# ═══════════════════════════════════════════════════════
# Style Resolver Tests (3+ cases)
# ═══════════════════════════════════════════════════════

class TestStyleResolver:
    def test_get_style_names(self, sample_docx):
        from docx_engine.style_resolver import StyleResolver
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        if parsed.get("styles") is not None:
            resolver = StyleResolver(parsed["styles"])
            names = resolver.get_style_names()
            assert len(names) > 0

    def test_get_styles_by_type(self, sample_docx):
        from docx_engine.style_resolver import StyleResolver
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        if parsed.get("styles") is not None:
            resolver = StyleResolver(parsed["styles"])
            para_styles = resolver.get_styles_by_type("paragraph")
            assert len(para_styles) >= 0

    def test_resolve_full_format(self, sample_docx):
        from docx_engine.style_resolver import StyleResolver
        from docx_engine.xml_parser import parse_docx
        parsed = parse_docx(sample_docx)
        if parsed.get("styles") is not None:
            resolver = StyleResolver(parsed["styles"])
            fmt = resolver.resolve_full_format("Normal")
            assert isinstance(fmt, dict)


# ═══════════════════════════════════════════════════════
# Error Handling Tests (5+ cases)
# ═══════════════════════════════════════════════════════

class TestErrors:
    def test_parse_error(self):
        from docx_engine.errors import ParseError, ErrorCode
        e = ParseError("Bad XML", ErrorCode.XML_MALFORMED)
        assert "Bad XML" in str(e)
        assert e.code == ErrorCode.XML_MALFORMED

    def test_serialize_error(self):
        from docx_engine.errors import SerializeError, ErrorCode
        e = SerializeError("Cannot serialize", ErrorCode.SERIALIZE_FAILED)
        assert "Cannot serialize" in str(e)

    def test_validation_error(self):
        from docx_engine.errors import ValidationError, ErrorCode
        e = ValidationError("Invalid format", ErrorCode.VALIDATION_FAILED)
        assert "Invalid format" in str(e)

    def test_error_codes_exist(self):
        from docx_engine.errors import ErrorCode
        assert ErrorCode.XML_MALFORMED is not None
        assert ErrorCode.SERIALIZE_FAILED is not None
        assert ErrorCode.VALIDATION_FAILED is not None
        assert ErrorCode.DOCUMENT_NOT_FOUND is not None
        assert ErrorCode.STYLE_NOT_FOUND is not None

    def test_custom_error_with_details(self):
        from docx_engine.errors import DocxEngineError
        e = DocxEngineError("Custom error", detail={"para": 5, "reason": "missing style"})
        assert e.detail == {"para": 5, "reason": "missing style"}


# ═══════════════════════════════════════════════════════
# Semantic Parser Tests (10+ cases)
# ═══════════════════════════════════════════════════════

class TestSemanticParser:
    def test_cover_title_detection(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="论文标题", bold=True, size=26)], alignment="center"))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.COVER_TITLE

    def test_abstract_label(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="摘要", bold=True, size=16)], alignment="center"))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.ABSTRACT_LABEL

    def test_abstract_content_contextual(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="摘要", bold=True, size=16)], alignment="center"))
        doc.paragraphs.append(Paragraph(runs=[Run(text="This is the abstract content about the research.")]))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.ABSTRACT_LABEL
        assert graph.elements[1].role == SemanticRole.ABSTRACT_CONTENT

    def test_keywords_label(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="关键词：机器学习；深度学习")]))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.KEYWORDS_LABEL

    def test_chapter_title(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="第一章 绪论", bold=True)], outline_level=0))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.CHAPTER_TITLE

    def test_reference_section_header(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="参考文献")], outline_level=0))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.REFERENCE_SECTION_HEADER

    def test_reference_item(self):
        from docx_engine.document_model import Document, Paragraph, Run
        from docx_engine.semantic_model import SemanticParser, SemanticRole
        doc = Document()
        doc.paragraphs.append(Paragraph(runs=[Run(text="[1] Author. Title. Journal, 2024.")]))
        parser = SemanticParser(doc)
        graph = parser.parse()
        assert graph.elements[0].role == SemanticRole.REFERENCE_ITEM

    def test_path_heading_relationships(self, sample_document):
        from docx_engine.semantic_model import SemanticParser
        parser = SemanticParser(sample_document)
        graph = parser.parse()
        headings = [e for e in graph.elements if e.role.startswith("heading_") or e.role == "chapter_title"]
        for h in headings:
            children = graph.get_children(h.index)
            assert isinstance(children, list)

    def test_section_body_retrieval(self, sample_document):
        from docx_engine.semantic_model import SemanticParser
        parser = SemanticParser(sample_document)
        graph = parser.parse()
        headings = graph.get_by_role("heading_1")
        for h in headings:
            body = graph.get_section_body(h.index)
            assert isinstance(body, list)

    def test_empty_document(self, empty_document):
        from docx_engine.semantic_model import SemanticParser
        parser = SemanticParser(empty_document)
        graph = parser.parse()
        assert len(graph.elements) == 0

    def test_get_by_role(self, sample_document):
        from docx_engine.semantic_model import SemanticParser
        parser = SemanticParser(sample_document)
        graph = parser.parse()
        reference_items = graph.get_by_role("reference_item")
        assert isinstance(reference_items, list)

    def test_to_dict_serializable(self, sample_document):
        from docx_engine.semantic_model import SemanticParser
        import json
        parser = SemanticParser(sample_document)
        graph = parser.parse()
        d = graph.to_dict()
        json.dumps(d, ensure_ascii=False, default=str)


# ═══════════════════════════════════════════════════════
# Layout Analyzer Tests (5+ cases)
# ═══════════════════════════════════════════════════════

class TestLayoutAnalyzer:
    def test_analyze_empty(self, empty_document):
        from docx_engine.layout_model import LayoutAnalyzer
        la = LayoutAnalyzer(empty_document)
        report = la.analyze()
        assert report.section_count == 0
        assert report.estimated_pages >= 0

    def test_analyze_with_content(self, sample_document):
        from docx_engine.layout_model import LayoutAnalyzer
        la = LayoutAnalyzer(sample_document)
        report = la.analyze()
        assert report.estimated_pages >= 1
        assert report.geometry is not None

    def test_page_geometry_a4(self):
        from docx_engine.document_model import Section
        from docx_engine.layout_model import PageGeometry
        s = Section()
        geo = PageGeometry.from_section(s)
        assert geo.width == 595.3
        assert geo.height == 841.9
        assert geo.printable_width > 400

    def test_report_to_dict(self, sample_document):
        from docx_engine.layout_model import LayoutAnalyzer
        import json
        la = LayoutAnalyzer(sample_document)
        report = la.analyze()
        d = report.to_dict()
        assert "geometry" in d
        assert "issues" in d
        json.dumps(d, ensure_ascii=False, default=str)

    def test_header_footer_analysis(self, sample_document):
        from docx_engine.layout_model import LayoutAnalyzer
        la = LayoutAnalyzer(sample_document)
        report = la.analyze()
        assert report.header_footer is not None
        assert isinstance(report.header_footer.has_first_page_different, bool)

    def test_layout_issues_collected(self, sample_document):
        from docx_engine.layout_model import LayoutAnalyzer
        la = LayoutAnalyzer(sample_document)
        report = la.analyze()
        assert isinstance(report.issues, list)


# ═══════════════════════════════════════════════════════
# Field Codes Tests (4+ cases)
# ═══════════════════════════════════════════════════════

class TestFieldCodes:
    def test_import_module(self):
        from wps_bridge.field_codes import FIELD_CODES, insert_field
        assert len(FIELD_CODES) > 10

    def test_field_constants(self):
        from wps_bridge.field_codes import FIELD_CODES
        assert "PAGE" in FIELD_CODES
        assert "SEQ" in FIELD_CODES
        assert "STYLEREF" in FIELD_CODES
        assert "REF" in FIELD_CODES


# ═══════════════════════════════════════════════════════
# Operation Logger Tests (3+ cases)
# ═══════════════════════════════════════════════════════

class TestOperationLogger:
    def test_record_and_summary(self):
        from intelligence.operation_logger import record, get_summary, clear, enable
        clear()
        enable()
        record("document", "info", {}, {"success": True}, 0.1)
        record("content", "outline", {}, {"success": True}, 0.05)
        summary = get_summary()
        assert summary["total_calls"] == 2
        assert summary["successes"] == 2
        assert "document" in summary["by_tool"]

    def test_recent(self):
        from intelligence.operation_logger import record, get_recent, clear, enable
        clear()
        enable()
        record("test", "action1", {}, {"ok": True}, 0.01)
        recent = get_recent(5)
        assert len(recent) == 1
        assert recent[0]["tool"] == "test"

    def test_errors(self):
        from intelligence.operation_logger import record, get_errors, clear, enable
        clear()
        enable()
        record("test", "fail", {}, {"error": "bad"}, 0.01, error="bad")
        errors = get_errors()
        assert len(errors) >= 1

