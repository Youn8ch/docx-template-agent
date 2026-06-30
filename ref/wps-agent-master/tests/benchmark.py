# -*- coding: utf-8 -*-
"""Performance benchmarks for wps-agent docx_engine."""
import os
import sys
import time
import tempfile
import shutil
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx_engine.document_model import Document, Paragraph, Run
from docx_engine.serializer import serialize_document_model
from docx_engine.semantic_model import SemanticParser
from docx_engine.layout_model import LayoutAnalyzer
from docx_engine.intelligence import DocumentAnalyzer
from docx_engine.formatter import Formatter


def _make_doc(n_paragraphs: int) -> Document:
    doc = Document()
    for i in range(1, n_paragraphs + 1):
        if i % 10 == 1:
            para = Paragraph(
                runs=[Run(text=f"Chapter {i//10 + 1} Section Title", bold=True, size=14)],
                outline_level=0,
            )
        elif i % 20 == 2:
            para = Paragraph(
                runs=[Run(text=f"Subsection {i//20 + 1} Topic", bold=True, size=12)],
                outline_level=1,
            )
        else:
            para = Paragraph(
                runs=[Run(text=f"This is paragraph number {i}. It contains sufficient text "
                      f"for realistic benchmarking of the document processing engine. "
                      f"The quick brown fox jumps over the lazy dog. "
                      f"人工智能技术正在深刻改变金融行业的面貌。")],
                first_line_indent=24,
            )
        doc.paragraphs.append(para)
    return doc


def benchmark(name: str, func, *args, **kwargs):
    start = time.perf_counter()
    result = func(*args, **kwargs)
    elapsed = time.perf_counter() - start
    print(f"  {name:40s} {elapsed*1000:8.1f} ms")
    return elapsed, result


def run(sizes=(100, 500, 1000)):
    print("=" * 60)
    print("WPS-Agent Performance Benchmarks")
    print("=" * 60)

    from docx import Document as PyDocx

    results = {}

    for size in sizes:
        print(f"\n{'─' * 60}")
        print(f"Document size: {size} paragraphs")
        print(f"{'─' * 60}")

        doc = _make_doc(size)

        # Save to disk — need a reference docx for serializer
        tmp = tempfile.mkdtemp()
        ref_path = os.path.join(tmp, "ref.docx")
        ref = PyDocx()
        ref.add_paragraph("Reference")
        ref.save(ref_path)

        path = os.path.join(tmp, f"bench_{size}.docx")
        serialize_document_model(doc, path, original_docx=ref_path)

        # Read back
        from docx_engine.serializer import load_and_build
        t, _ = benchmark("Load & Parse", load_and_build, path)
        doc = load_and_build(path)

        # Semantic
        t, graph = benchmark("Semantic Parse", SemanticParser(doc).parse)

        # Intelligence
        analyzer = DocumentAnalyzer(doc)
        t, _ = benchmark("Document Analysis", analyzer.analyze_formatting_quality)

        # Layout
        t, _ = benchmark("Layout Analysis", LayoutAnalyzer(doc).analyze)

        # Format
        formatter = Formatter(doc)
        t, _ = benchmark("Auto Format (thesis)", formatter.auto_format, "thesis")

        # Full pipeline
        from intelligence.format_intelligence import auto_enhance
        t, result = benchmark("Full Auto-Enhance", auto_enhance, path, None, None, path)
        print(f"    → Score: {result.get('final_score', 0)}, Verdict: {result.get('verdict', 'N/A')}")

        results[size] = {
            "ms_total": t * 1000,
            "ms_per_paragraph": (t / size) * 1000,
        }

        shutil.rmtree(tmp)

    print(f"\n{'=' * 60}")
    print("Summary")
    print(f"{'=' * 60}")
    for size, data in results.items():
        print(f"  {size:5d} paragraphs → Full Enhance: {data['ms_total']:8.1f} ms ({data['ms_per_paragraph']:6.2f} ms/para)")

    return results


if __name__ == "__main__":
    run()
