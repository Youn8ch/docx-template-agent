# -*- coding: utf-8 -*-
"""
Verification script for the refactored WPS-Agent document engine.
Tests core modules without requiring WPS to be running.
"""
import sys
import os
import tempfile
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

def test_imports():
    """Test that all core modules can be imported."""
    print("=== Phase 1: Module Imports ===")
    core_modules = [
        "docx_engine.errors",
        "docx_engine.xml_parser",
        "docx_engine.document_model",
        "docx_engine.style_resolver",
        "docx_engine.serializer",
        "docx_engine.intelligence",
        "docx_engine.formatter",
        "offline.docx_builder",
    ]
    com_modules = [
        "wps_bridge.utils",
        "wps_bridge.app",
        "wps_bridge.com_client",
        "wps_bridge.commands",
    ]

    for mod in core_modules:
        try:
            __import__(mod)
            print(f"  [OK] {mod}")
        except Exception as e:
            print(f"  [FAIL] {mod}: {e}")
            return False

    for mod in com_modules:
        try:
            __import__(mod)
            print(f"  [OK] {mod}")
        except ImportError as e:
            if "pythoncom" in str(e) or "pywin32" in str(e).lower() or "No module named" in str(e):
                print(f"  [SKIP] {mod}: {e}")
            else:
                print(f"  [FAIL] {mod}: {e}")
                return False
        except Exception as e:
            print(f"  [FAIL] {mod}: {e}")
            return False

    print()
    return True


def test_xml_parser():
    """Test ZIP unpack/repack and XML parsing."""
    print("=== Phase 2: XML Parser ===")
    from docx_engine.xml_parser import unpack_docx, pack_docx, parse_docx, get_paragraphs, get_paragraph_text
    from docx import Document

    # Create a test docx
    tmpdir = tempfile.mkdtemp()
    test_path = os.path.join(tmpdir, "test.docx")
    doc = Document()
    doc.add_heading("Test Heading 1", level=1)
    doc.add_paragraph("This is a test paragraph with some text.")
    doc.add_paragraph("Second paragraph with bold text.")
    doc.save(test_path)
    print(f"  [OK] Created test docx: {test_path}")

    # Unpack
    unpacked = unpack_docx(test_path)
    print(f"  [OK] Unpacked to: {unpacked}")
    assert os.path.exists(os.path.join(unpacked, "word", "document.xml"))

    # Parse
    parsed = parse_docx(test_path)
    print(f"  [OK] Parsed document.xml")
    assert parsed["document"] is not None

    # Read paragraphs
    paras = get_paragraphs(parsed["document"])
    print(f"  [OK] Found {len(paras)} paragraphs")
    assert len(paras) >= 3  # heading + 2 paragraphs

    texts = [get_paragraph_text(p) for p in paras]
    print(f"  [OK] Paragraph texts: {texts}")

    # Repack
    output_path = os.path.join(tmpdir, "output.docx")
    pack_docx(unpacked, output_path)
    print(f"  [OK] Repacked to: {output_path}")
    assert os.path.exists(output_path)

    # Cleanup
    import shutil
    shutil.rmtree(tmpdir)
    print()
    return True


def test_document_model():
    """Test Document DOM operations."""
    print("=== Phase 3: Document Model ===")
    from docx_engine.document_model import Document, Paragraph, Run

    doc = Document()

    # Add paragraphs
    p1 = Paragraph(runs=[Run(text="Hello ", bold=True), Run(text="World", italic=True)])
    p1.style_id = "Heading1"
    p1.outline_level = 0
    doc.paragraphs.append(p1)

    p2 = Paragraph(runs=[Run(text="This is body text.", font="宋体", size=12)])
    p2.alignment = "justify"
    p2.first_line_indent = 24
    doc.paragraphs.append(p2)

    # Test properties
    assert doc.text == "Hello World\nThis is body text."
    print(f"  [OK] Document text: {repr(doc.text)}")

    assert p1.is_heading()
    assert p1.heading_level() == 1
    print(f"  [OK] Heading detection works")

    assert len(p1.runs) == 2
    assert p1.runs[0].bold
    assert p1.runs[1].italic
    print(f"  [OK] Run-level properties preserved")

    # Test search
    indices = doc.find_paragraphs_by_text("body")
    assert indices == [2]
    print(f"  [OK] Text search works: found at indices {indices}")

    # Test statistics
    stats = doc.get_statistics()
    assert stats["paragraph_count"] == 2
    print(f"  [OK] Statistics: {stats}")

    # Test heading structure
    headings = doc.get_heading_structure()
    assert len(headings) == 1
    print(f"  [OK] Heading structure: {headings}")

    print()
    return True


def test_serializer():
    """Test XML serialization round-trip."""
    print("=== Phase 4: Serializer Round-Trip ===")
    from docx_engine.serializer import load_and_build, serialize_document_model
    from docx import Document

    tmpdir = tempfile.mkdtemp()
    test_path = os.path.join(tmpdir, "test.docx")
    output_path = os.path.join(tmpdir, "output.docx")

    # Create test docx with formatting
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Bold ").bold = True
    p.add_run("Italic ").italic = True
    p.add_run("Normal")
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Body text here.")
    doc.save(test_path)

    # Load into model
    model = load_and_build(test_path)
    print(f"  [OK] Loaded model: {len(model.paragraphs)} paragraphs")
    assert len(model.paragraphs) >= 3

    # Check run-level precision
    first_para = model.paragraphs[0]
    assert len(first_para.runs) == 3
    assert first_para.runs[0].bold
    assert first_para.runs[1].italic
    assert not first_para.runs[2].bold
    print(f"  [OK] Run-level format preserved through round-trip")

    # Check heading
    heading = None
    for p in model.paragraphs:
        if p.is_heading():
            heading = p
            break
    assert heading is not None
    assert heading.heading_level() == 1
    print(f"  [OK] Heading preserved: level={heading.heading_level()}, text={heading.text}")

    # Serialize back
    serialize_document_model(model, output_path, original_docx=test_path)
    assert os.path.exists(output_path)
    print(f"  [OK] Serialized back to: {output_path}")

    # Re-read and verify
    model2 = load_and_build(output_path)
    assert len(model2.paragraphs) >= 3
    print(f"  [OK] Re-read model: {len(model2.paragraphs)} paragraphs")

    # Cleanup
    import shutil
    shutil.rmtree(tmpdir)
    print()
    return True


def test_intelligence():
    """Test document intelligence and analysis."""
    print("=== Phase 5: Document Intelligence ===")
    from docx_engine.document_model import Document, Paragraph, Run
    from docx_engine.intelligence import DocumentAnalyzer

    doc = Document()
    doc.paragraphs.append(Paragraph(runs=[Run(text="Abstract", bold=True, size=16)], alignment="center"))
    doc.paragraphs.append(Paragraph(runs=[Run(text="This is the abstract content.")]))
    doc.paragraphs.append(Paragraph(runs=[Run(text="1 Introduction", bold=True, size=14)], outline_level=0))
    doc.paragraphs.append(Paragraph(runs=[Run(text="This is the introduction.")]))
    doc.paragraphs.append(Paragraph(runs=[Run(text="2 Methods", bold=True, size=14)], outline_level=0))
    doc.paragraphs.append(Paragraph(runs=[Run(text="Methodology description.")]))
    doc.paragraphs.append(Paragraph(runs=[Run(text="References")], outline_level=0))
    doc.paragraphs.append(Paragraph(runs=[Run(text="[1] Author, Title, Journal, 2024.")]))

    analyzer = DocumentAnalyzer(doc)

    # Detect type
    doc_type = analyzer.detect_document_type()
    print(f"  [OK] Detected type: {doc_type}")

    # Detect roles
    roles = analyzer.detect_paragraph_roles()
    role_names = [r["role"] for r in roles]
    print(f"  [OK] Detected roles: {role_names}")
    assert "heading1" in role_names
    assert "abstract_label" in role_names or "abstract_content" in role_names

    # Analyze quality
    quality = analyzer.analyze_formatting_quality()
    print(f"  [OK] Quality score: {quality.get('score', 0)}")

    # Get outline
    outline = analyzer.get_document_outline()
    print(f"  [OK] Outline: {[n['text'] for n in outline]}")

    print()
    return True


def test_formatter():
    """Test professional formatter."""
    print("=== Phase 6: Professional Formatter ===")
    from docx_engine.document_model import Document, Paragraph, Run
    from docx_engine.formatter import Formatter

    doc = Document()
    doc.paragraphs.append(Paragraph(runs=[Run(text="Chapter 1 Introduction", bold=True)], outline_level=0))
    doc.paragraphs.append(Paragraph(runs=[Run(text="This is the first paragraph of the introduction.")]))
    doc.paragraphs.append(Paragraph(runs=[Run(text="Section 1.1 Background", bold=True)], outline_level=1))
    doc.paragraphs.append(Paragraph(runs=[Run(text="Background information here.")]))

    formatter = Formatter(doc)

    # Auto-format as thesis
    result = formatter.auto_format("thesis")
    print(f"  [OK] Auto-formatted thesis: {result['change_count']} changes")

    # Check heading format
    p1 = doc.paragraphs[0]
    assert p1.alignment == "center"
    assert p1.space_before == 24
    print(f"  [OK] Heading 1 formatted: alignment={p1.alignment}, space_before={p1.space_before}")

    # Check body format
    p2 = doc.paragraphs[1]
    assert p2.alignment == "justify"
    print(f"  [OK] Body formatted: alignment={p2.alignment}")

    # Test numbering
    result = formatter.add_multi_level_numbering()
    print(f"  [OK] Numbering applied: {result['change_count']} changes")
    assert "第1章" in doc.paragraphs[0].text or "1." in doc.paragraphs[0].text
    print(f"  [OK] Heading text: {doc.paragraphs[0].text}")

    print()
    return True


def test_offline_builder():
    """Test offline builder read/write."""
    print("=== Phase 7: Offline Builder ===")
    from offline.docx_builder import OfflineDocxBuilder
    from docx import Document

    tmpdir = tempfile.mkdtemp()
    test_path = os.path.join(tmpdir, "test.docx")

    # Create test file
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("Paragraph one.")
    doc.add_paragraph("Paragraph two.")
    doc.save(test_path)

    # Load with builder
    builder = OfflineDocxBuilder()
    builder.load(test_path)
    print(f"  [OK] Loaded: {len(builder.document.paragraphs)} paragraphs")

    # Analyze
    analysis = builder.analyze()
    print(f"  [OK] Analysis: type={analysis['document_type']}, stats={analysis['statistics']}")

    # Replace text
    count = builder.replace_text("Paragraph", "Section")
    print(f"  [OK] Replaced {count} occurrences")

    # Save
    output_path = os.path.join(tmpdir, "output.docx")
    builder.save(output_path)
    assert os.path.exists(output_path)
    print(f"  [OK] Saved to: {output_path}")

    # Verify
    doc2 = Document(output_path)
    texts = [p.text for p in doc2.paragraphs]
    print(f"  [OK] Output texts: {texts}")
    assert any("Section" in t for t in texts)

    # Cleanup
    import shutil
    shutil.rmtree(tmpdir)
    print()
    return True


def test_style_resolver():
    """Test style resolver with real styles.xml."""
    print("=== Phase 8: Style Resolver ===")
    from docx_engine.style_resolver import StyleResolver
    from docx import Document
    from docx_engine.xml_parser import parse_docx

    tmpdir = tempfile.mkdtemp()
    test_path = os.path.join(tmpdir, "test.docx")

    doc = Document()
    doc.add_heading("Title", level=1)
    doc.save(test_path)

    parsed = parse_docx(test_path)
    if parsed.get("styles") is not None:
        resolver = StyleResolver(parsed["styles"])
        names = resolver.get_style_names()
        print(f"  [OK] Found {len(names)} styles")
        assert len(names) > 0

        heading_styles = resolver.get_styles_by_type("paragraph")
        print(f"  [OK] Paragraph styles: {len(heading_styles)}")
    else:
        print("  [INFO] No styles.xml in test docx (normal for python-docx default)")

    import shutil
    shutil.rmtree(tmpdir)
    print()
    return True


def run_all_tests():
    """Run all verification tests."""
    print("=" * 60)
    print("WPS-Agent v2 Document Engine Verification")
    print("=" * 60)
    print()

    tests = [
        test_imports,
        test_xml_parser,
        test_document_model,
        test_serializer,
        test_intelligence,
        test_formatter,
        test_offline_builder,
        test_style_resolver,
    ]

    passed = 0
    failed = 0

    for test in tests:
        try:
            if test():
                passed += 1
            else:
                failed += 1
        except Exception as e:
            print(f"  [FAIL] {test.__name__}: {e}")
            import traceback
            traceback.print_exc()
            failed += 1

    print("=" * 60)
    print(f"Results: {passed} passed, {failed} failed out of {len(tests)} tests")
    print("=" * 60)
    return failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
