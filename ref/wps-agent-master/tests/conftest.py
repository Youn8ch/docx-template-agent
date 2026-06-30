# -*- coding: utf-8 -*-
"""Shared fixtures for wps-agent tests."""
import os
import sys
import tempfile
import shutil
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))


@pytest.fixture
def sample_run():
    from docx_engine.document_model import Run
    return Run(text="Hello", font="宋体", size=12, bold=True)


@pytest.fixture
def sample_paragraph():
    from docx_engine.document_model import Paragraph, Run
    return Paragraph(
        runs=[
            Run(text="Bold text ", bold=True, size=14),
            Run(text="italic text ", italic=True),
            Run(text="normal text", size=12),
        ],
        alignment="justify",
        first_line_indent=24,
        space_before=12,
        space_after=6,
    )


@pytest.fixture
def sample_heading_paragraph():
    from docx_engine.document_model import Paragraph, Run
    return Paragraph(
        runs=[Run(text="Chapter 1 Introduction", bold=True, size=16)],
        style_id="Heading1",
        outline_level=0,
        alignment="center",
        space_before=24,
        space_after=12,
    )


@pytest.fixture
def sample_table():
    from docx_engine.document_model import Table, Cell, Paragraph, Run
    rows = []
    for r in range(3):
        row = []
        for c in range(2):
            cell = Cell(
                paragraphs=[Paragraph(runs=[Run(text=f"R{r}C{c}")])]
            )
            row.append(cell)
        rows.append(row)
    return Table(rows=rows, column_widths=[200, 300])


@pytest.fixture
def sample_document():
    from docx_engine.document_model import Document, Paragraph, Run
    doc = Document()
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="My Thesis Title", bold=True, size=22)],
        style_id="Title",
        alignment="center",
        outline_level=0,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Abstract", bold=True, size=16)],
        alignment="center",
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="This is the abstract content describing the research.")],
        alignment="justify",
        first_line_indent=24,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="1 Introduction", bold=True, size=14)],
        outline_level=0,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Introduction body text here.")],
        first_line_indent=24,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="2 Methods", bold=True, size=14)],
        outline_level=0,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Methodology description with formula E=mc².")],
        first_line_indent=24,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="References", bold=True)],
        outline_level=0,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="[1] Author A. Title. Journal, 2024.")],
    ))
    return doc


@pytest.fixture
def empty_document():
    from docx_engine.document_model import Document
    return Document()


@pytest.fixture
def temp_dir():
    d = tempfile.mkdtemp()
    yield d
    try:
        shutil.rmtree(d)
    except Exception:
        pass


@pytest.fixture
def sample_docx(temp_dir):
    from docx import Document as PyDocx
    path = os.path.join(temp_dir, "test.docx")
    doc = PyDocx()
    doc.add_heading("Test Document", level=1)
    p = doc.add_paragraph()
    p.add_run("Bold ").bold = True
    p.add_run("Italic ").italic = True
    p.add_run("Normal")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("Body text for section 1.")
    doc.save(path)
    return path


@pytest.fixture
def sample_docx_with_table(temp_dir):
    from docx import Document as PyDocx
    path = os.path.join(temp_dir, "table_test.docx")
    doc = PyDocx()
    doc.add_heading("Table Test", level=1)
    table = doc.add_table(rows=3, cols=3)
    for r, row_data in enumerate([["A1", "B1", "C1"], ["A2", "B2", "C2"], ["A3", "B3", "C3"]]):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val
    doc.save(path)
    return path


@pytest.fixture
def sample_contract_doc():
    from docx_engine.document_model import Document, Paragraph, Run
    doc = Document()
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="SERVICE AGREEMENT", bold=True, size=18)],
        alignment="center",
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="This Agreement is made on January 1, 2024.")],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="1. Definitions", bold=True, size=14)],
        outline_level=0,
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text='"Party A" means the Service Provider.')],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="2. Terms and Conditions", bold=True, size=14)],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="The parties hereby agree to the following terms.")],
    ))
    return doc


@pytest.fixture
def sample_resume_doc():
    from docx_engine.document_model import Document, Paragraph, Run
    doc = Document()
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="John Doe", bold=True, size=20)],
        alignment="center",
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Email: john@example.com | Phone: 555-0100")],
        alignment="center",
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Education", bold=True, size=14)],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="BS Computer Science, University of Example, 2020")],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Experience", bold=True, size=14)],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Senior Developer at Tech Corp, 2020-2024")],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Skills", bold=True, size=14)],
    ))
    doc.paragraphs.append(Paragraph(
        runs=[Run(text="Python, Java, Docker, AWS, Machine Learning")],
    ))
    return doc
