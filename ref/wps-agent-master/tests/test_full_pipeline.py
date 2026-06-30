# -*- coding: utf-8 -*-
"""End-to-end pipeline test: .docx → semantic parse → format → quality → verify."""
import os
import sys
import tempfile
import shutil
import pytest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))


@pytest.fixture
def sample_thesis_docx():
    """Create a realistic thesis-style .docx for pipeline testing."""
    from docx import Document as PyDocx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmp = tempfile.mkdtemp()
    path = os.path.join(tmp, "thesis.docx")
    doc = PyDocx()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Deep Learning Applications in Financial Risk Management").bold = True
    title.runs[0].font.size = 26 * 12700  # 26pt

    abstract_label = doc.add_paragraph()
    abstract_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_label.add_run("Abstract").bold = True

    doc.add_paragraph(
        "This paper investigates the application of deep learning techniques "
        "to financial risk management, focusing on credit risk assessment "
        "and market risk prediction."
    )

    doc.add_paragraph("Keywords: deep learning, risk management, neural networks")

    # Chapters
    doc.add_heading("1 Introduction", level=1)
    doc.add_paragraph(
        "Financial risk management has evolved significantly with the advent "
        "of artificial intelligence and machine learning technologies."
    )
    doc.add_paragraph(
        "Traditional risk models such as Value at Risk (VaR) and Monte Carlo "
        "simulation have limitations in capturing complex nonlinear patterns."
    )

    doc.add_heading("2 Literature Review", level=1)
    doc.add_paragraph(
        "Previous research has explored various neural network architectures "
        "for financial time series prediction."
    )

    doc.add_heading("3 Methodology", level=1)
    doc.add_paragraph(
        "We propose a hybrid model combining LSTM networks with attention "
        "mechanisms for credit risk prediction."
    )

    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Smith J. Neural Networks in Finance. J Finance, 2023.")
    doc.add_paragraph("[2] Brown A. Deep Learning for Risk. Risk Management, 2024.")

    doc.save(path)
    yield path
    shutil.rmtree(tmp)


class TestFullPipeline:
    def test_auto_enhance_thesis(self, sample_thesis_docx):
        from intelligence.format_intelligence import auto_enhance
        result = auto_enhance(sample_thesis_docx)
        assert "error" not in result
        assert "stages" in result
        assert result["stages"]["parse"]["status"] == "ok"
        assert result["stages"]["semantic"]["status"] == "ok"
        assert result["stages"]["format"]["status"] == "ok"
        assert result["final_score"] >= 0

    def test_auto_enhance_document_type(self, sample_thesis_docx):
        from intelligence.format_intelligence import auto_enhance
        result = auto_enhance(sample_thesis_docx)
        doc_type = result["stages"]["semantic"].get("document_type", "")
        assert doc_type in ("thesis", "paper", "report")

    def test_auto_enhance_creates_output(self, sample_thesis_docx):
        import tempfile
        from intelligence.format_intelligence import auto_enhance
        tmp = tempfile.mkdtemp()
        output = os.path.join(tmp, "enhanced.docx")
        result = auto_enhance(sample_thesis_docx, output_path=output)
        assert os.path.exists(output)
        shutil.rmtree(tmp)

    def test_auto_enhance_heading_format(self, sample_thesis_docx):
        from intelligence.format_intelligence import auto_enhance
        from docx import Document as PyDocx

        result = auto_enhance(sample_thesis_docx)
        doc = PyDocx(sample_thesis_docx)
        headings_found = False
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                headings_found = True
                break
        assert headings_found

    def test_semantic_pipeline_integration(self, sample_thesis_docx):
        from docx_engine import parse_docx, build_document_model
        from docx_engine.semantic_model import SemanticParser, SemanticRole

        parsed = parse_docx(sample_thesis_docx)
        doc = build_document_model(parsed)
        parser = SemanticParser(doc)
        graph = parser.parse()

        roles = [e.role for e in graph.elements]
        assert SemanticRole.ABSTRACT_LABEL in roles or SemanticRole.ABSTRACT_CONTENT in roles
        assert SemanticRole.KEYWORDS_LABEL in roles
        assert SemanticRole.REFERENCE_ITEM in roles

    def test_format_pipeline_integration(self, sample_thesis_docx):
        from docx_engine import parse_docx, build_document_model
        from docx_engine.formatter import Formatter

        parsed = parse_docx(sample_thesis_docx)
        doc = build_document_model(parsed)
        formatter = Formatter(doc)
        result = formatter.auto_format("thesis")
        assert result["applied"] is True
        assert result["change_count"] > 0

    def test_layout_pipeline_integration(self, sample_thesis_docx):
        from docx_engine import parse_docx, build_document_model
        from docx_engine.layout_model import LayoutAnalyzer

        parsed = parse_docx(sample_thesis_docx)
        doc = build_document_model(parsed)
        la = LayoutAnalyzer(doc)
        report = la.analyze()
        assert report.estimated_pages >= 1
        assert len(report.to_dict()) > 0
