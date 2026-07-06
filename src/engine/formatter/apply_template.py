"""Apply whitelisted formatting operations to a new docx file."""

from __future__ import annotations

import os
import warnings
from pathlib import Path
from uuid import uuid4
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.engine.checker.operation_validator import validate_operations
from src.engine.checker.style_checker import check_styles
from src.engine.formatter.font_utils import first_line_indent_chars_value, set_run_font
from src.engine.model.document_model import DocumentModel, ParagraphInfo
from src.engine.model.operation_model import (
    ExecutionReport,
    Operation,
    OperationResult,
    StyleCheckReport,
    is_safe_operation,
)
from src.engine.parser.docx_parser import parse_docx
from src.engine.parser.structure_detector import detect_structure
from src.engine.safety.content_integrity import (
    compare_content_snapshots,
    content_snapshot,
)


ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HEADING_ROLES = {"title", "heading_1", "heading_2", "heading_3"}
FIRST_LINE_INDENT_ATTRS = ("w:firstLineChars", "w:firstLine", "w:hanging")


def _resolve_output_path(input_path: Path, output_path: str | Path) -> Path:
    target = Path(output_path)
    if target.suffix.lower() == ".docx":
        resolved = target.resolve()
    else:
        resolved = (target / f"{input_path.stem}_formatted.docx").resolve()

    if resolved == input_path.resolve():
        raise ValueError("output path cannot overwrite the source docx")
    return resolved


def _paragraph_by_index(document: Any, index: int) -> Any:
    if index < 1 or index > len(document.paragraphs):
        raise IndexError(f"paragraph index out of range: {index}")
    return document.paragraphs[index - 1]


def _table_by_index(document: Any, index: int) -> Any:
    if index < 1 or index > len(document.tables):
        raise IndexError(f"table index out of range: {index}")
    return document.tables[index - 1]


def _set_alignment(paragraph: Any, alignment: str | None) -> None:
    if alignment is None:
        return
    key = str(alignment).strip().lower()
    if key not in ALIGNMENT_MAP:
        raise ValueError(f"unsupported paragraph alignment: {alignment}")
    paragraph.alignment = ALIGNMENT_MAP[key]


def _get_or_add_indent(paragraph: Any) -> Any:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.ind
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    return ind


def _clear_indent_attrs(paragraph: Any, attrs: tuple[str, ...] = FIRST_LINE_INDENT_ATTRS) -> None:
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.ind is None:
        return
    ind = p_pr.ind
    for attr in attrs:
        ind.attrib.pop(qn(attr), None)


def _set_left_indent_zero(paragraph: Any) -> None:
    ind = _get_or_add_indent(paragraph)
    ind.set(qn("w:left"), "0")


def _set_first_line_indent_chars(paragraph: Any, properties: dict[str, Any]) -> None:
    value = first_line_indent_chars_value(properties)
    if value is None:
        return

    ind = _get_or_add_indent(paragraph)
    for attr in ("w:firstLine", "w:hanging", "w:hangingChars"):
        ind.attrib.pop(qn(attr), None)
    ind.set(qn("w:firstLineChars"), str(value))
    ind.set(qn("w:left"), "0")


def _set_paragraph_format(paragraph: Any, properties: dict[str, Any], role: str | None) -> None:
    paragraph_format = paragraph.paragraph_format
    _set_alignment(paragraph, properties.get("alignment"))

    if "line_spacing" in properties and properties["line_spacing"] is not None:
        paragraph_format.line_spacing = float(properties["line_spacing"])
    if "space_before" in properties and properties["space_before"] is not None:
        paragraph_format.space_before = Pt(float(properties["space_before"]))
    if "space_after" in properties and properties["space_after"] is not None:
        paragraph_format.space_after = Pt(float(properties["space_after"]))

    if role == "body":
        _set_first_line_indent_chars(paragraph, properties)
    elif role in HEADING_ROLES:
        if "first_line_indent_chars" in properties and properties["first_line_indent_chars"] is not None:
            _set_first_line_indent_chars(paragraph, properties)
        else:
            _clear_indent_attrs(paragraph)
            _set_left_indent_zero(paragraph)


def _format_runs(paragraphs: list[Any], properties: dict[str, Any]) -> None:
    for paragraph in paragraphs:
        for run in paragraph.runs:
            set_run_font(run, properties)


def _apply_paragraph_style(document: Any, operation: Operation) -> None:
    for index in operation.target_indices:
        paragraph = _paragraph_by_index(document, index)
        _format_runs([paragraph], operation.properties)
        _set_paragraph_format(paragraph, operation.properties, operation.role)


def _apply_table_style(
    document: Any,
    operation: Operation,
    header_table_indices: set[int] | None = None,
) -> None:
    header_table_indices = header_table_indices or set()
    for index in operation.target_indices:
        table = _table_by_index(document, index)
        rows = table.rows[1:] if index in header_table_indices else table.rows
        for row in rows:
            for cell in row.cells:
                _format_runs(list(cell.paragraphs), operation.properties)
                for paragraph in cell.paragraphs:
                    _set_alignment(paragraph, operation.properties.get("alignment"))
                    _clear_indent_attrs(paragraph)


def _apply_table_header_style(document: Any, operation: Operation) -> None:
    for index in operation.target_indices:
        table = _table_by_index(document, index)
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            _format_runs(list(cell.paragraphs), operation.properties)
            for paragraph in cell.paragraphs:
                _set_alignment(paragraph, operation.properties.get("alignment"))
                _clear_indent_attrs(paragraph)


def _apply_operation(document: Any, operation: Operation) -> None:
    if operation.action == "apply_paragraph_style":
        _apply_paragraph_style(document, operation)
    elif operation.action == "apply_table_style":
        _apply_table_style(document, operation)
    elif operation.action == "apply_table_header_style":
        _apply_table_header_style(document, operation)
    else:
        raise ValueError(f"unsupported operation action: {operation.action}")


def _execute_operations(document: Any, operations: list[Operation]) -> list[OperationResult]:
    results: list[OperationResult] = []
    header_table_indices = {
        index
        for operation in operations
        if operation.action == "apply_table_header_style"
        for index in operation.target_indices
    }

    for operation in operations:
        if not is_safe_operation(operation):
            results.append(
                OperationResult(
                    operation_id=operation.operation_id,
                    status="failed",
                    message=f"operation rejected by whitelist: {operation.action}",
                )
            )
            continue

        try:
            if operation.action == "apply_table_style":
                _apply_table_style(document, operation, header_table_indices)
            else:
                _apply_operation(document, operation)
        except Exception as exc:
            results.append(
                OperationResult(
                    operation_id=operation.operation_id,
                    status="failed",
                    message=f"operation failed: {exc}",
                )
            )
            continue

        results.append(
            OperationResult(
                operation_id=operation.operation_id,
                status="success",
                message=f"operation applied: {operation.action}",
            )
        )

    return results


def _temp_output_path(target: Path) -> Path:
    return target.with_name(f".{target.stem}.{uuid4().hex}.tmp{target.suffix}")


def _remove_temp_file(temp_path: Path) -> None:
    try:
        if temp_path.exists():
            temp_path.unlink()
    except OSError:
        pass


def _temp_report_fields(temp_path: Path, retain_failed_temp: bool) -> dict[str, Any]:
    if retain_failed_temp and temp_path.exists():
        return {"temp_output_path": str(temp_path), "temp_file_retained": True}
    _remove_temp_file(temp_path)
    return {"temp_output_path": None, "temp_file_retained": False}


def _role_map(document: DocumentModel) -> dict[int, str]:
    return {paragraph.index: paragraph.role for paragraph in document.paragraphs}


def _reuse_roles(
    parsed_after: DocumentModel,
    document_before: DocumentModel,
) -> tuple[DocumentModel | None, list[str]]:
    errors: list[str] = []
    before_indices = [paragraph.index for paragraph in document_before.paragraphs]
    after_indices = [paragraph.index for paragraph in parsed_after.paragraphs]
    if before_indices != after_indices:
        errors.append("paragraph index structure changed")
        return None, errors
    if document_before.paragraph_count != parsed_after.paragraph_count:
        errors.append("paragraph count changed before role reuse")
        return None, errors

    roles = _role_map(document_before)
    paragraphs: list[ParagraphInfo] = [
        paragraph.model_copy(update={"role": roles[paragraph.index]})
        for paragraph in parsed_after.paragraphs
    ]
    return parsed_after.model_copy(update={"paragraphs": paragraphs}), errors


def apply_operations(
    input_path: str | Path,
    output_path: str | Path,
    operations: list[Operation],
) -> list[dict[str, str]]:
    warnings.warn(
        "apply_operations is deprecated; use apply_operations_transactional instead",
        DeprecationWarning,
        stacklevel=2,
    )
    source = Path(input_path).resolve()
    if not source.exists():
        raise FileNotFoundError(f"input docx not found: {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError(f"input must be a .docx file: {source}")

    target = _resolve_output_path(source, output_path)
    document_before = detect_structure(parse_docx(source))
    roles = {
        str(operation.role): {}
        for operation in operations
        if operation.role is not None
    }
    roles.update({"table": {}, "table_header": {}})
    template = {"template_id": "deprecated_apply_operations", "rules": roles}
    validation_errors = validate_operations(operations, document_before, template)
    if validation_errors:
        raise ValueError("; ".join(validation_errors))

    report = apply_operations_transactional(
        source,
        target,
        operations,
        template=template,
        document_before=document_before,
        retain_failed_temp=False,
    )
    if report.status != "success":
        errors = report.validation_errors or report.integrity_errors or report.recheck_errors
        raise RuntimeError(f"failed to apply operations safely: {report.status}: {'; '.join(errors)}")
    return report.execution_results


def apply_operations_transactional(
    input_path: str | Path,
    output_path: str | Path,
    operations: list[Operation],
    *,
    template: dict[str, Any],
    document_before: DocumentModel,
    report_before: StyleCheckReport | None = None,
    retain_failed_temp: bool = False,
    overwrite_output: bool = True,
) -> ExecutionReport:
    source = Path(input_path).resolve()
    if not source.exists():
        return ExecutionReport(status="fatal", recheck_errors=[f"input docx not found: {source}"])
    if source.suffix.lower() != ".docx":
        return ExecutionReport(status="fatal", recheck_errors=[f"input must be a .docx file: {source}"])

    try:
        target = _resolve_output_path(source, output_path)
    except Exception as exc:
        return ExecutionReport(status="fatal", recheck_errors=[str(exc)])

    try:
        target.parent.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        return ExecutionReport(
            status="fatal",
            expected_output_path=str(target),
            recheck_errors=[f"{exc.__class__.__name__}: {exc}"],
        )
    output_existed_before = target.exists()
    if output_existed_before and not overwrite_output:
        before_report = report_before or check_styles(document_before, template)
        return ExecutionReport(
            status="validation_failed",
            expected_output_path=str(target),
            issues_before=before_report.issue_count,
            operations_before=len(operations),
            validation_errors=[f"output path already exists and overwrite_output is false: {target}"],
            integrity_before=content_snapshot(document_before),
            output_existed_before=True,
            output_overwritten=False,
        )

    before_report = report_before or check_styles(document_before, template)
    before_snapshot = content_snapshot(document_before)
    validation_errors = validate_operations(operations, document_before, template)
    if validation_errors:
        return ExecutionReport(
            status="validation_failed",
            expected_output_path=str(target),
            issues_before=before_report.issue_count,
            operations_before=len(operations),
            validation_errors=validation_errors,
            integrity_before=before_snapshot,
            output_existed_before=output_existed_before,
        )

    temp_path = _temp_output_path(target)
    try:
        document = Document(str(source))
        operation_results = _execute_operations(document, operations)
        failed = [result for result in operation_results if result.status == "failed"]
        if failed:
            return ExecutionReport(
                status="execution_failed",
                expected_output_path=str(target),
                issues_before=before_report.issue_count,
                operations_before=len(operations),
                execution_results=[result.model_dump() for result in operation_results],
                integrity_before=before_snapshot,
                output_existed_before=output_existed_before,
                **_temp_report_fields(temp_path, retain_failed_temp),
            )

        document.save(str(temp_path))
        parsed_after = parse_docx(temp_path)
        after_snapshot = content_snapshot(parsed_after)
        integrity_errors = compare_content_snapshots(before_snapshot, after_snapshot)
        parsed_after_with_roles, role_errors = _reuse_roles(parsed_after, document_before)
        recheck_errors = role_errors
        if document_before.table_count != parsed_after.table_count:
            recheck_errors.append("table count changed before recheck")

        if integrity_errors or recheck_errors or parsed_after_with_roles is None:
            return ExecutionReport(
                status="content_integrity_failed",
                expected_output_path=str(target),
                issues_before=before_report.issue_count,
                issues_after=None,
                operations_before=len(operations),
                operations_after=None,
                execution_results=[result.model_dump() for result in operation_results],
                integrity_before=before_snapshot,
                integrity_after=after_snapshot,
                integrity_errors=integrity_errors,
                recheck_errors=recheck_errors,
                output_existed_before=output_existed_before,
                **_temp_report_fields(temp_path, retain_failed_temp),
            )

        after_report = check_styles(parsed_after_with_roles, template)
        if after_report.issue_count > 0 or after_report.operation_count > 0:
            recheck_errors = [
                "style recheck still has "
                f"{after_report.issue_count} issue(s) and {after_report.operation_count} operation(s)"
            ]
            return ExecutionReport(
                status="validation_failed",
                expected_output_path=str(target),
                issues_before=before_report.issue_count,
                issues_after=after_report.issue_count,
                operations_before=len(operations),
                operations_after=after_report.operation_count,
                execution_results=[result.model_dump() for result in operation_results],
                integrity_before=before_snapshot,
                integrity_after=after_snapshot,
                recheck_errors=recheck_errors,
                output_existed_before=output_existed_before,
                **_temp_report_fields(temp_path, retain_failed_temp),
            )

        os.replace(temp_path, target)
        return ExecutionReport(
            status="success",
            output_path=str(target),
            expected_output_path=str(target),
            temp_output_path=None,
            temp_file_retained=False,
            output_existed_before=output_existed_before,
            output_overwritten=output_existed_before,
            issues_before=before_report.issue_count,
            issues_after=after_report.issue_count,
            operations_before=len(operations),
            operations_after=after_report.operation_count,
            execution_results=[result.model_dump() for result in operation_results],
            integrity_before=before_snapshot,
            integrity_after=after_snapshot,
        )
    except Exception as exc:
        temp_fields = _temp_report_fields(temp_path, retain_failed_temp)
        return ExecutionReport(
            status="fatal",
            expected_output_path=str(target),
            output_existed_before=output_existed_before,
            issues_before=before_report.issue_count,
            operations_before=len(operations),
            integrity_before=before_snapshot,
            recheck_errors=[f"{exc.__class__.__name__}: {exc}"],
            **temp_fields,
        )
