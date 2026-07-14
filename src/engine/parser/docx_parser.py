"""Read-only docx parser.

The parser extracts document structure and observed formatting into
DocumentModel. It never saves the document or mutates the source file.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.opc.exceptions import PackageNotFoundError

from src.engine.model.document_model import (
    DocumentModel,
    ParagraphInfo,
    RunInfo,
    TableCellInfo,
    TableInfo,
)


def _style_name(obj: Any) -> str | None:
    style = getattr(obj, "style", None)
    return getattr(style, "name", None) if style is not None else None


def _alignment_name(alignment: Any) -> str | None:
    if alignment is None:
        return None
    name = getattr(alignment, "name", None)
    if name:
        return str(name).lower()
    return str(alignment).lower()


def _length_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    pt = getattr(value, "pt", None)
    if pt is not None:
        return float(pt)
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _line_spacing(value: Any) -> float | None:
    return _length_to_pt(value)


def _first_line_indent_chars(paragraph: Any) -> float | None:
    p_pr = paragraph._p.pPr
    ind = getattr(p_pr, "ind", None) if p_pr is not None else None
    if ind is None:
        return None
    value = ind.get(qn("w:firstLineChars"))
    if value is None:
        return None
    try:
        return float(value) / 100
    except (TypeError, ValueError):
        return None


def _indent_length(paragraph: Any, attr: str) -> float | None:
    p_pr = paragraph._p.pPr
    ind = getattr(p_pr, "ind", None) if p_pr is not None else None
    if ind is None:
        return None
    value = ind.get(qn(attr))
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _indent_chars(paragraph: Any, attr: str) -> float | None:
    value = _indent_length(paragraph, attr)
    if value is None:
        return None
    return value / 100


def _run_font_name(run: Any) -> str | None:
    font_name = getattr(run.font, "name", None)
    if font_name:
        return font_name

    r_pr = getattr(run._element, "rPr", None)
    r_fonts = getattr(r_pr, "rFonts", None) if r_pr is not None else None
    if r_fonts is None:
        return None
    return r_fonts.get(qn("w:eastAsia")) or r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi"))


def _run_font_size(run: Any) -> float | None:
    size = getattr(run.font, "size", None)
    return _length_to_pt(size)


def _paragraph_runs(paragraph: Any) -> list[RunInfo]:
    runs: list[RunInfo] = []
    for index, run in enumerate(paragraph.runs, start=1):
        runs.append(
            RunInfo(
                index=index,
                text=run.text or "",
                font_name=_run_font_name(run),
                font_size=_run_font_size(run),
                bold=run.font.bold,
            )
        )
    return runs


def _paragraph_info(paragraph: Any, index: int) -> ParagraphInfo:
    paragraph_format = paragraph.paragraph_format
    runs = _paragraph_runs(paragraph)
    return ParagraphInfo(
        index=index,
        text=paragraph.text or "",
        style_name=_style_name(paragraph),
        alignment=_alignment_name(paragraph.alignment),
        font_names=[run.font_name for run in runs],
        font_sizes=[run.font_size for run in runs],
        bold_values=[run.bold for run in runs],
        line_spacing=_line_spacing(paragraph_format.line_spacing),
        space_before=_length_to_pt(paragraph_format.space_before),
        space_after=_length_to_pt(paragraph_format.space_after),
        first_line_indent=_length_to_pt(paragraph_format.first_line_indent),
        first_line=_indent_length(paragraph, "w:firstLine"),
        first_line_indent_chars=_first_line_indent_chars(paragraph),
        hanging=_indent_length(paragraph, "w:hanging"),
        hanging_chars=_indent_chars(paragraph, "w:hangingChars"),
        runs=runs,
    )


def _table_cell_info(cell: Any, row_index: int, col_index: int) -> TableCellInfo:
    font_names: list[str | None] = []
    font_sizes: list[float | None] = []
    bold_values: list[bool | None] = []
    style_name: str | None = None
    alignment: str | None = None

    for paragraph in cell.paragraphs:
        if style_name is None:
            style_name = _style_name(paragraph)
        if alignment is None:
            alignment = _alignment_name(paragraph.alignment)
        for run in _paragraph_runs(paragraph):
            font_names.append(run.font_name)
            font_sizes.append(run.font_size)
            bold_values.append(run.bold)

    return TableCellInfo(
        row_index=row_index,
        col_index=col_index,
        text=cell.text or "",
        paragraph_texts=[paragraph.text or "" for paragraph in cell.paragraphs],
        style_name=style_name,
        alignment=alignment,
        font_names=font_names,
        font_sizes=font_sizes,
        bold_values=bold_values,
    )


def _table_info(table: Any, index: int) -> TableInfo:
    rows = len(table.rows)
    cols = max((len(row.cells) for row in table.rows), default=0)
    cells: list[TableCellInfo] = []

    for row_index, row in enumerate(table.rows, start=1):
        for col_index, cell in enumerate(row.cells, start=1):
            cells.append(_table_cell_info(cell, row_index, col_index))

    return TableInfo(index=index, rows=rows, cols=cols, cells=cells)


def _used_styles(paragraphs: list[ParagraphInfo], tables: list[TableInfo]) -> list[str]:
    styles: set[str] = set()
    for paragraph in paragraphs:
        if paragraph.style_name:
            styles.add(paragraph.style_name)
    for table in tables:
        for cell in table.cells:
            if cell.style_name:
                styles.add(cell.style_name)
    return sorted(styles)


def parse_docx(path: str | Path, include_facts: bool = False) -> DocumentModel:
    docx_path = Path(path)
    if not docx_path.exists():
        raise FileNotFoundError(f"input docx not found: {docx_path}")
    if not docx_path.is_file():
        raise ValueError(f"input path is not a file: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"input must be a .docx file: {docx_path}")

    try:
        document = Document(str(docx_path))
    except PackageNotFoundError as exc:
        raise ValueError(f"invalid or unreadable docx file: {docx_path}") from exc
    except Exception as exc:
        raise RuntimeError(f"failed to read docx file: {docx_path}") from exc

    paragraphs = [
        _paragraph_info(paragraph, index)
        for index, paragraph in enumerate(document.paragraphs, start=1)
    ]
    tables = [_table_info(table, index) for index, table in enumerate(document.tables, start=1)]

    model = DocumentModel(
        filepath=docx_path,
        paragraph_count=len(paragraphs),
        table_count=len(tables),
        paragraphs=paragraphs,
        tables=tables,
        styles=_used_styles(paragraphs, tables),
    )
    if include_facts:
        from src.engine.parser.facts_extractor import extract_document_facts

        model = model.model_copy(update={"facts": extract_document_facts(docx_path, model)})
    return model
